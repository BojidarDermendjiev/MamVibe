"""
MamVibe — Project Overview Document Generator
Generates a comprehensive .docx report for the MamVibe platform.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import copy

# ──────────────────────────────────────────────────────────
# Brand palette  (warm, family-friendly)
# ──────────────────────────────────────────────────────────
PURPLE       = RGBColor(0x7C, 0x3A, 0xED)   # primary violet
PURPLE_DARK  = RGBColor(0x5B, 0x21, 0xB6)   # headings
PURPLE_LIGHT = RGBColor(0xED, 0xE9, 0xFE)   # table header fill
WARM_GRAY    = RGBColor(0x6B, 0x72, 0x80)   # body text
ACCENT_PINK  = RGBColor(0xF4, 0x72, 0xB6)   # accent
ACCENT_TEAL  = RGBColor(0x14, 0xB8, 0xA6)   # secondary accent
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
DARK         = RGBColor(0x1F, 0x2A, 0x3C)   # near-black for cover text
LIGHT_GRAY   = RGBColor(0xF9, 0xFA, 0xFB)   # alt row fill

# ──────────────────────────────────────────────────────────
# Low-level XML helpers
# ──────────────────────────────────────────────────────────

def rgb_to_hex(rgb: RGBColor) -> str:
    """Convert RGBColor to 6-char hex string."""
    return str(rgb)  # RGBColor.__str__ returns e.g. '5B21B6'


def set_cell_bg(cell, rgb: RGBColor):
    """Fill a table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = rgb_to_hex(rgb)
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Add borders to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),   kwargs.get("val",   "single"))
        tag.set(qn("w:sz"),    kwargs.get("sz",    "4"))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color", "E5E7EB"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def set_run_font(run, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def page_break(doc):
    doc.add_page_break()


def add_horizontal_rule(doc):
    """Thin purple rule between sections."""
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr= OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "7C3AED")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(12)


# ──────────────────────────────────────────────────────────
# Styled paragraph helpers
# ──────────────────────────────────────────────────────────

def add_cover_title(doc, text, size=44):
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(size)
    run.font.color.rgb = WHITE
    set_run_font(run)
    return p


def add_cover_subtitle(doc, text, size=18, color=None):
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color or PURPLE_LIGHT
    set_run_font(run)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = PURPLE_DARK
    set_run_font(run)
    add_horizontal_rule(doc)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = PURPLE
    set_run_font(run)
    return p


def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK
    set_run_font(run)
    return p


def body(doc, text, italic=False, color=None):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.italic    = italic
    run.font.color.rgb = color or WARM_GRAY
    set_run_font(run)
    return p


def bullet(doc, text, level=0):
    p   = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = WARM_GRAY
    set_run_font(run)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(3)
    return p


def kv_row(doc, key, value):
    """Key: Value inline paragraph."""
    p    = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    rkey = p.add_run(f"{key}:  ")
    rkey.bold = True
    rkey.font.size = Pt(10.5)
    rkey.font.color.rgb = DARK
    set_run_font(rkey)
    rval = p.add_run(value)
    rval.font.size = Pt(10.5)
    rval.font.color.rgb = WARM_GRAY
    set_run_font(rval)
    return p


# ──────────────────────────────────────────────────────────
# Table helper
# ──────────────────────────────────────────────────────────

def add_table(doc, headers, rows, col_widths=None, alt_rows=True):
    """
    Styled table with purple header row.
    headers: list[str]
    rows:    list[list[str]]
    col_widths: list[float] in inches (optional)
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style     = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, PURPLE_DARK)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p    = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        set_run_font(run)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg  = PURPLE_LIGHT if (alt_rows and r_idx % 2 == 0) else WHITE
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p    = cell.paragraphs[0]
            run  = p.add_run(cell_text)
            run.font.size = Pt(10)
            run.font.color.rgb = WARM_GRAY
            set_run_font(run)

    # Column widths
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    doc.add_paragraph()  # spacing after table
    return table


# ──────────────────────────────────────────────────────────
# Cover page (purple gradient simulation via shaded paragraphs)
# ──────────────────────────────────────────────────────────

def build_cover(doc):
    # Simulate a dark purple cover by using a full-width shaded paragraph block
    # We'll use a 1-cell table spanning the page as a colour block

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, PURPLE_DARK)
    cell.width = Inches(6.5)

    # Clear default empty paragraph inside cell
    cell.paragraphs[0].clear()

    def cover_para(text, size, color=WHITE, bold=False, space_before=0, space_after=6, center=True):
        p   = cell.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        run = p.add_run(text)
        run.bold = bold
        run.font.size  = Pt(size)
        run.font.color.rgb = color
        set_run_font(run)
        return p

    # Top breathing room
    cover_para("", 8, space_before=60, space_after=0)

    # Logo text / brand mark
    cover_para("✦  MamVibe  ✦", 13, color=ACCENT_PINK, bold=False, space_after=4)

    # Main title
    cover_para("Platform Overview", 38, bold=True, space_before=8, space_after=8)

    # Divider line via run character
    cover_para("─" * 45, 10, color=PURPLE_LIGHT, space_after=16)

    # Tagline
    cover_para(
        "The modern marketplace where families connect,\n"
        "donate, and shop for baby & child essentials.",
        13, color=PURPLE_LIGHT, space_after=32
    )

    # Meta info block
    cover_para(f"Report Date:  {datetime.date.today().strftime('%B %d, %Y')}", 10, color=PURPLE_LIGHT, space_after=4)
    cover_para("Version:  1.0", 10, color=PURPLE_LIGHT, space_after=4)
    cover_para("Confidential — Internal & Investor Use Only", 9, color=ACCENT_PINK, space_after=60)

    doc.add_paragraph()  # space after cover table
    page_break(doc)


# ──────────────────────────────────────────────────────────
# Section builders
# ──────────────────────────────────────────────────────────

def section_executive_summary(doc):
    heading1(doc, "1. Executive Summary")

    body(doc,
        "MamVibe is a full-stack digital marketplace designed for families with young children. "
        "It provides a safe, community-driven platform where parents can donate, buy, or sell "
        "second-hand baby and children's items — from clothes and strollers to toys and accessories. "
        "The platform combines a rich e-commerce experience with real-time chat, AI-powered tools, "
        "integrated logistics, and a vibrant community hub."
    )

    body(doc,
        "Built with enterprise-grade technology (.NET 8 on the backend, React 19 on the frontend), "
        "MamVibe is production-ready and designed to scale. It integrates with leading Bulgarian "
        "shipping couriers, Stripe for card payments, an internal escrow wallet system, Cloudflare "
        "for security and storage, and AI providers (Anthropic Claude, Groq) for smart listing "
        "assistance and content moderation."
    )

    body(doc,
        "MamVibe targets Bulgarian families as its primary market with bilingual support "
        "(English and Bulgarian), Bulgarian fiscal compliance (TakeANap integration, 20 % VAT), "
        "and integrations with the four most popular local couriers. The platform is architected "
        "to expand into additional markets with minimal engineering effort."
    )

    heading2(doc, "At a Glance")
    add_table(doc,
        headers=["Dimension", "Detail"],
        rows=[
            ["Platform type",      "Peer-to-peer marketplace (C2C) + community hub"],
            ["Primary market",     "Bulgaria (bilingual EN / BG)"],
            ["Target audience",    "Families with children aged 0-16"],
            ["Listing types",      "Donate (free) & Sell (priced)"],
            ["Payment methods",    "Card (Stripe), Wallet escrow, On-spot, Cash on Delivery"],
            ["Shipping couriers",  "Econt, Speedy, Box Now, Pigeon Express"],
            ["Real-time features", "SignalR chat, typing indicators, push notifications"],
            ["AI features",        "Listing suggestions, content moderation, AI chat bot"],
            ["Backend",            ".NET 8 / ASP.NET Core — Clean Architecture"],
            ["Frontend",           "React 19 + TypeScript + Vite + Tailwind CSS"],
            ["Database",           "PostgreSQL 18 + PgBouncer + Redis"],
            ["Infrastructure",     "Docker Compose, Nginx, Grafana, Prometheus, Loki"],
        ],
        col_widths=[2.2, 4.3]
    )


def section_product_overview(doc):
    heading1(doc, "2. Product Overview")

    body(doc,
        "MamVibe was founded on the belief that baby items — often used for only a few months — "
        "deserve a second life. Rather than ending up in landfills, prams, clothes, and toys "
        "can circulate within a trusted community of parents. MamVibe makes this exchange "
        "simple, safe, and socially rewarding."
    )

    heading2(doc, "2.1  Mission")
    body(doc,
        "To build the most trusted and convenient platform for families to exchange children's "
        "items, fostering sustainability, community, and financial savings for parents across Bulgaria."
    )

    heading2(doc, "2.2  Core User Journeys")

    heading3(doc, "Seller / Donor Journey")
    for step in [
        "Register and set up a profile (individual, couple, or family).",
        "Create a listing: upload up to 5 photos, set a title, description, category, condition, age group, clothing/shoe size, and price (or mark as free donation).",
        "AI assistant analyses the photo and pre-fills listing details to save time.",
        "Listing goes live after passing AI moderation and (optionally) admin review.",
        "Receive purchase requests or direct Stripe checkout from buyers.",
        "Accept request → shipment is created automatically with the chosen courier.",
        "Funds are held in escrow and released to the seller's wallet upon confirmed delivery.",
        "Seller receives a rating from the buyer.",
    ]:
        bullet(doc, step)

    heading3(doc, "Buyer Journey")
    for step in [
        "Browse items with filters: category, listing type, age group, size, condition, price range.",
        "Like and save items; create saved searches with email alerts for new matches.",
        "Send a purchase request or buy instantly via Stripe card.",
        "Choose shipping: courier office, home address, or parcel locker.",
        "Track the shipment in real time from the dashboard.",
        "Confirm receipt → funds are released to the seller.",
        "Leave a star rating and review for the seller.",
    ]:
        bullet(doc, step)

    heading3(doc, "Community Journey")
    for step in [
        "Follow other users and browse their listings.",
        "Chat in real time with sellers or buyers.",
        "Discover and review local doctors and child-friendly places (playgrounds, cafés, parks).",
        "Submit platform feedback (bug reports, feature ideas).",
        "Access bilingual content (EN / BG) based on personal preference.",
    ]:
        bullet(doc, step)


def section_features(doc):
    heading1(doc, "3. Feature Set")

    heading2(doc, "3.1  Marketplace Core")
    for f in [
        "Item listings with up to 5 photos, title, description, category, condition, age group, clothing and shoe sizes.",
        "Two listing types: Donate (free) and Sell (priced).",
        "Advanced browse & search: filtering by category, listing type, age group, size, condition, price range, and keyword.",
        "Pagination and multiple sort modes (newest, price, popularity).",
        "View count and like/favourite tracking per item.",
        "Bump (re-list) items to surface them again after 7 days.",
        "Bundle creation: group 2-10 items at a combined discounted price.",
        "Purchase request workflow: buyer requests → seller accepts or declines.",
        "Reserve status: mark items as reserved without fully selling them.",
        "Holiday mode: hide all seller listings temporarily.",
        "Stale item detection: flag listings active for 30 + days.",
        "AI-powered listing assistant: photo analysis → auto-filled title, description, price suggestion, and category.",
        "AI content moderation: automatic screening of descriptions with confidence scores and admin fallback.",
    ]:
        bullet(doc, f)

    heading2(doc, "3.2  Payments")
    add_table(doc,
        headers=["Payment Method", "Provider", "Flow"],
        rows=[
            ["Card Checkout",    "Stripe",          "Stripe Checkout session → webhook confirmation → funds to escrow wallet"],
            ["Wallet",          "Internal escrow", "Buyer tops up wallet via Stripe PaymentIntent → funds held → released on delivery"],
            ["On-spot",         "Internal",        "Seller manually records cash transaction; platform tracks without movement of funds"],
            ["Cash on Delivery","Courier COD",     "Courier collects cash at doorstep; platform syncs status on delivery confirmation"],
            ["Booking",         "Internal",        "Free donation reservation; no financial transaction required"],
        ],
        col_widths=[1.8, 1.8, 3.0]
    )

    heading2(doc, "3.3  Shipping & Logistics")
    body(doc,
        "MamVibe integrates with all four major Bulgarian courier networks. A pluggable provider "
        "architecture means new couriers can be added with minimal engineering effort."
    )
    add_table(doc,
        headers=["Courier", "Delivery Types", "Features"],
        rows=[
            ["Econt Express",   "Office / Home / Locker", "Full API: price calc, label, tracking, COD, insurance"],
            ["Speedy",          "Office / Home",           "Full API: price calc, label, tracking, COD"],
            ["Box Now",         "Locker / Office",         "Parcel locker network, price calc, tracking"],
            ["Pigeon Express",  "Home / Office",           "API integration, tracking, COD support"],
        ],
        col_widths=[1.6, 2.0, 3.0]
    )
    body(doc, "Shared shipping features:")
    for f in [
        "Shipping price calculation per courier, delivery type, and parcel weight.",
        "Automated waybill and label generation; PDF download from the dashboard.",
        "Real-time tracking sync: Pending → Created → Picked Up → In Transit → Out for Delivery → Delivered.",
        "COD (cash on delivery) and insurance options.",
        "Office and locker lookup API for address-free pickup.",
        "Admin shipment overview with bulk status management.",
    ]:
        bullet(doc, f)

    heading2(doc, "3.4  Real-time Messaging & Notifications")
    for f in [
        "WebSocket chat powered by SignalR at /hubs/chat.",
        "Typing indicators and read receipts.",
        "User presence tracking — online/offline badges throughout the UI.",
        "Unread message badges with counts in the navigation bar.",
        "Offline message email notification: if the recipient is not connected, an email is dispatched.",
        "AI chat bot: a special system user powered by Anthropic Claude or Groq that answers questions about listings and platform usage.",
        "Server-push notifications for purchase requests, offer updates, new follows, and shipment status changes.",
        "Expo push notifications for mobile clients.",
    ]:
        bullet(doc, f)

    heading2(doc, "3.5  Community & Social")
    for f in [
        "User profiles: display name, avatar (photo or illustrated), bio, profile type (Male / Female / Family).",
        "Follow / follower system with social feed integration.",
        "Post-purchase ratings (1-5 stars) with written comments; shown on public profiles.",
        "Doctor reviews: crowdsourced reviews of paediatricians, dentists, and specialists — admin-moderated before publishing.",
        "Child-friendly places: community directory of parks, playgrounds, family restaurants, and more — with age-range filters.",
        "Feedback system: users submit bug reports, feature requests, or general feedback with a star rating.",
    ]:
        bullet(doc, f)

    heading2(doc, "3.6  Financial Compliance")
    for f in [
        "Electronic bills (e-bills) auto-generated for every completed payment.",
        "Unique, idempotent e-bill numbering: MV-YYYY-{first 8 hex chars of payment ID}.",
        "E-bill email delivery with rate limiting (3 resend/minute per user).",
        "TakeANap fiscal receipt integration: HMAC-SHA256 signed receipts compliant with Bulgarian tax authority requirements.",
        "VAT calculation at the standard Bulgarian rate of 20 %.",
    ]:
        bullet(doc, f)

    heading2(doc, "3.7  Admin Panel")
    for f in [
        "Dashboard with live statistics: active items, pending payments, in-transit shipments, recent feedback.",
        "User management: view profiles, block / unblock accounts.",
        "Item moderation: approve or reject listings; view AI moderation scores.",
        "Shipping oversight: full shipment list with status filters and courier breakdowns.",
        "Community moderation: approve doctor reviews and child-friendly places.",
        "Audit log viewer: filterable record of all admin actions with timestamps and change diffs.",
    ]:
        bullet(doc, f)

    heading2(doc, "3.8  Offers & Negotiation")
    for f in [
        "Buyers can submit a price offer below the listed price.",
        "Sellers can counter-offer with a different price.",
        "Offer status lifecycle: Pending → Accepted / Declined.",
        "Accepted offers automatically transition into a purchase request.",
    ]:
        bullet(doc, f)

    heading2(doc, "3.9  Saved Searches & Alerts")
    for f in [
        "Save any search configuration as a named alert.",
        "Filters: category, listing type, keyword, age group, size, condition, max price.",
        "Automated email alert when a new item matches a saved search.",
        "Manage saved searches from the user dashboard.",
    ]:
        bullet(doc, f)

    heading2(doc, "3.10  Security & Compliance")
    for f in [
        "Cloudflare Turnstile CAPTCHA on all authentication forms (registration, login, password reset).",
        "JWT access tokens (short-lived) + HttpOnly refresh token cookies.",
        "Google OAuth single sign-on.",
        "ASP.NET Identity password policy: minimum 8 characters, uppercase, lowercase, digit, and special character.",
        "Identity lockout: 5 failed attempts triggers a 5-minute lockout.",
        "Global rate limiting (200 req/min), tighter limits on auth (30/min) and upload (20/min) endpoints.",
        "GDPR data export endpoint: users can download all their personal data.",
        "Blocked user middleware: suspended accounts are rejected on every request.",
        "Full HTTP security header suite: CSP, X-Frame-Options, HSTS, Referrer-Policy, Permissions-Policy, CORP, COOP.",
    ]:
        bullet(doc, f)


def section_tech_stack(doc):
    heading1(doc, "4. Technology Stack")

    heading2(doc, "4.1  Backend")
    add_table(doc,
        headers=["Component", "Technology", "Purpose"],
        rows=[
            ["Framework",         ".NET 8 / ASP.NET Core",        "Web API host, dependency injection, middleware"],
            ["Architecture",      "Clean Architecture (4 layers)", "Domain → Application → Infrastructure → WebApi"],
            ["Database",          "PostgreSQL 18",                 "Primary relational data store"],
            ["ORM",               "Entity Framework Core 9",       "Code-first migrations, LINQ queries"],
            ["Connection pool",   "PgBouncer (transaction mode)",  "500 max clients, 20 pool size — reduces DB connections"],
            ["Cache / PubSub",    "Redis 7",                       "Distributed cache, SignalR backplane"],
            ["Real-time",         "ASP.NET Core SignalR",          "WebSocket chat and server-push notifications"],
            ["Authentication",    "ASP.NET Identity + JWT",        "Token-based auth with refresh tokens"],
            ["OAuth",             "Google OAuth 2.0",              "Single sign-on via Google accounts"],
            ["Validation",        "FluentValidation",              "Request DTO validation in Application layer"],
            ["Mapping",           "AutoMapper",                    "Entity ↔ DTO mapping profiles"],
            ["Logging",           "Serilog",                       "Structured logging with sink to Loki"],
            ["Observability",     "OpenTelemetry",                 "Distributed traces (OTLP export), metrics"],
            ["Background jobs",   "IHostedService / Channel<T>",   "n8n dispatcher, tracking sync, scheduled alerts"],
            ["API versioning",    "Asp.Versioning",                "URL-based versioning (/api/v1/)"],
        ],
        col_widths=[1.8, 2.2, 2.6]
    )

    heading2(doc, "4.2  Frontend")
    add_table(doc,
        headers=["Component", "Technology", "Purpose"],
        rows=[
            ["Framework",        "React 19 + TypeScript",      "Component-based UI with full type safety"],
            ["Build tool",       "Vite",                       "Fast HMR dev server, optimised production bundle"],
            ["Styling",          "Tailwind CSS",               "Utility-first responsive design"],
            ["State management", "Zustand",                    "Lightweight global state with localStorage persistence"],
            ["HTTP client",      "Axios",                      "API requests with interceptors (auth, refresh, error)"],
            ["Real-time",        "@microsoft/signalr",         "WebSocket connection to /hubs/chat"],
            ["Routing",          "React Router v6",            "Client-side routing with protected routes"],
            ["i18n",             "react-i18next",              "English and Bulgarian UI translations"],
            ["Notifications",    "react-hot-toast",            "In-app toast notifications"],
            ["Date handling",    "date-fns",                   "Locale-aware date formatting"],
            ["Path alias",       "@  →  src/",                 "Configured in vite.config.ts for clean imports"],
            ["Dev proxy",        "Vite proxy",                 "Proxies /api, /hubs, /uploads to localhost:5038"],
        ],
        col_widths=[1.8, 2.2, 2.6]
    )

    heading2(doc, "4.3  Infrastructure")
    add_table(doc,
        headers=["Service", "Technology", "Role"],
        rows=[
            ["Container orchestration", "Docker + Docker Compose",    "Full-stack local & production deployment"],
            ["Reverse proxy",           "Nginx",                      "Static file serving, API proxying, SSL termination"],
            ["Metrics",                 "Prometheus",                 "Scrapes /metrics from API; stores time-series data"],
            ["Logs",                    "Loki (31-day retention)",    "Aggregates structured logs from all services"],
            ["Dashboards & alerts",     "Grafana",                    "Visualises metrics + logs; sends alert notifications"],
            ["Photo storage",           "Cloudflare R2 (S3-compat.)", "Object storage for item photos and avatars"],
            ["Bot protection",          "Cloudflare Turnstile",       "Invisible CAPTCHA on auth forms"],
            ["Secrets management",      "Doppler (optional)",         "Centralised secret injection at deploy time"],
            ["Data Protection keys",    "Docker volume",              "Persists ASP.NET Data Protection keys across restarts"],
        ],
        col_widths=[2.0, 2.2, 2.4]
    )

    heading2(doc, "4.4  Third-party Integrations")
    add_table(doc,
        headers=["Integration", "Purpose"],
        rows=[
            ["Stripe",               "Card checkout sessions, PaymentIntent (wallet top-up), webhooks"],
            ["Econt Express API",    "Shipping price calc, label generation, tracking, COD"],
            ["Speedy API",           "Shipping price calc, label generation, tracking, COD"],
            ["Box Now API",          "Parcel locker shipping, tracking"],
            ["Pigeon Express API",   "Courier delivery, tracking, COD"],
            ["TakeANap",             "Bulgarian fiscal receipt generation (HMAC-SHA256 signed)"],
            ["Anthropic Claude",     "AI listing assistant, content moderation, chat bot (vision + text)"],
            ["Groq",                 "Alternative AI chat provider (low-latency inference)"],
            ["n8n",                  "Automation workflows: email notifications, seller reports, alerts"],
            ["SMTP",                 "Transactional email: password reset, e-bills, offline notifications"],
            ["Google OAuth",         "Single sign-on for users without email/password credentials"],
            ["Cloudflare R2",        "S3-compatible object storage for user-uploaded photos"],
            ["Expo Push",            "Push notifications to native mobile clients"],
        ],
        col_widths=[2.0, 4.5]
    )


def section_architecture(doc):
    heading1(doc, "5. Architecture")

    heading2(doc, "5.1  Backend — Clean Architecture")
    body(doc,
        "The backend is structured as four separate .NET projects with strictly enforced "
        "dependency flow: outer layers may depend on inner layers, never the reverse."
    )
    add_table(doc,
        headers=["Layer", "Project", "Contains"],
        rows=[
            ["Domain",         "MomVibe.Domain",         "Entities, enums, constants. Zero external dependencies."],
            ["Application",    "MomVibe.Application",    "DTOs, service interfaces, FluentValidation validators, AutoMapper profiles."],
            ["Infrastructure", "MomVibe.Infrastructure", "EF Core DbContext + migrations, external API clients (Stripe, couriers, AI), service implementations."],
            ["Presentation",   "MomVibe.WebApi",         "Controllers, SignalR hubs, middleware pipeline, DI wiring (StartUp.cs)."],
        ],
        col_widths=[1.4, 2.0, 3.2]
    )

    heading2(doc, "5.2  Key Architectural Patterns")

    heading3(doc, "N8n Webhook Dispatcher (Fire-and-forget)")
    body(doc,
        "Business events (payment completed, shipment delivered, user registered, etc.) are "
        "enqueued into a bounded Channel<T> with capacity 500. A BackgroundService drains "
        "the queue asynchronously and POSTs payloads to the configured n8n webhook URL. "
        "This decouples the request path from external automation latency."
    )

    heading3(doc, "Shipping Provider Factory")
    body(doc,
        "A CourierProviderFactory resolves the correct ICourierProvider implementation "
        "(Econt, Speedy, Box Now, Pigeon Express) based on the enum value on the Shipment "
        "entity. Adding a new courier requires only a new class implementing the interface "
        "and a registration in the DI container — no changes to existing code."
    )

    heading3(doc, "Payment Idempotency")
    body(doc,
        "An optional Idempotency-Key request header is stored in the Payment entity. "
        "A unique database index prevents duplicate Payment rows. The key is forwarded "
        "to Stripe to prevent duplicate charges even under retry conditions."
    )

    heading3(doc, "Output Cache for High-Traffic Endpoints")
    body(doc,
        "The item browse endpoint is cached for 30 seconds using ASP.NET Output Cache. "
        "The categories endpoint is cached for 1 hour with tag-based invalidation. "
        "Redis is the backing store, ensuring consistency across multiple API instances."
    )

    heading3(doc, "SignalR with Redis Backplane")
    body(doc,
        "All real-time events go through SignalR. The Redis backplane ensures that "
        "messages and notifications are broadcast correctly even when multiple API "
        "container instances are running behind the Nginx load balancer."
    )

    heading3(doc, "E-Bill Idempotency")
    body(doc,
        "E-bill numbers are assigned exactly once. The EBillNumber column is checked "
        "before assignment: if already set (e.g., from a Stripe webhook replay), no "
        "second number is generated. Format: MV-{YEAR}-{first 8 hex chars of payment ID}."
    )

    heading2(doc, "5.3  Database Schema — Key Entities")
    add_table(doc,
        headers=["Entity", "Key Fields", "Notes"],
        rows=[
            ["ApplicationUser",     "DisplayName, ProfileType, AvatarUrl, IsBlocked, Bio, LanguagePreference, RevolutTag, ExpoPushToken, IsOnHoliday", "Extends ASP.NET Identity"],
            ["Item",                "Title, Description, CategoryId, ListingType, AgeGroup, Size, Price, UserId, IsActive, IsReserved, IsSold, ViewCount, Condition, AiModerationStatus", "Core marketplace entity"],
            ["Payment",             "ItemId, BuyerId, SellerId, Amount, PaymentMethod, PaymentStatus, StripeSessionId, EBillNumber, IdempotencyKey", "Covers all payment methods"],
            ["Shipment",            "PaymentId, CourierProvider, DeliveryType, Status, TrackingNumber, WaybillId, ShippingPrice, IsCod, CodAmount, IsInsured", "Linked 1:1 to Payment"],
            ["Message",             "SenderId, ReceiverId, Content, IsRead", "Chat message; read in SignalR hub"],
            ["PurchaseRequest",     "ItemId, BuyerId, SellerId, Status", "Pending / Accepted / Declined"],
            ["Offer",               "ItemId, BuyerId, SellerId, OfferedPrice, CounterPrice, Status", "Price negotiation entity"],
            ["UserRating",          "RaterId, RatedUserId, PurchaseRequestId, Rating, Comment", "Post-transaction review"],
            ["DoctorReview",        "DoctorName, Specialization, City, Rating, IsApproved", "Community-moderated"],
            ["ChildFriendlyPlace",  "Name, PlaceType, City, AgeFromMonths, AgeToMonths, IsApproved", "Community directory"],
            ["SavedSearch",         "UserId, Name, Filters (category, type, size, price…)", "Alert trigger on new items"],
            ["AuditLog",            "UserId, Action, EntityType, EntityId, Changes, CreatedAt", "Admin action history"],
        ],
        col_widths=[1.7, 3.0, 1.9]
    )

    heading2(doc, "5.4  Frontend Structure")
    add_table(doc,
        headers=["Directory", "Contents"],
        rows=[
            ["src/api/",        "Axios client modules per domain (auth, items, payments, shipping, messages…)"],
            ["src/store/",      "Zustand stores: authStore, itemStore, chatStore, notificationStore…"],
            ["src/contexts/",   "React contexts: SignalR connection, auth provider"],
            ["src/pages/",      "One file per route: HomePage, BrowsePage, ItemDetailPage, DashboardPage, ChatPage, AdminPage…"],
            ["src/components/", "Reusable UI components: Navbar, ItemCard, PaymentModal, ShipmentTracker, ChatWindow…"],
            ["src/layouts/",    "MainLayout (nav + footer), AuthLayout (centered), AdminLayout (sidebar)"],
            ["src/locales/",    "en.json and bg.json translation files"],
            ["src/hooks/",      "Custom hooks: useAuth, useSignalR, useInfiniteScroll…"],
        ],
        col_widths=[1.8, 4.8]
    )


def section_api_reference(doc):
    heading1(doc, "6. API Reference Summary")

    body(doc,
        "All endpoints are versioned under /api/v1/. Protected endpoints require "
        "a valid JWT Bearer token. Admin endpoints require the Admin role."
    )

    heading2(doc, "6.1  Authentication  (/api/v1/auth)")
    add_table(doc,
        headers=["Method", "Path", "Auth", "Description"],
        rows=[
            ["POST", "/register",        "Public",     "Register with email, password, and Turnstile token"],
            ["POST", "/login",           "Public",     "Email/password login; returns access + refresh token"],
            ["POST", "/refresh",         "Public",     "Exchange refresh token for new access token"],
            ["POST", "/google-login",    "Public",     "Google OAuth sign-in / sign-up"],
            ["GET",  "/me",              "Authorized", "Fetch current user's profile"],
            ["POST", "/revoke",          "Authorized", "Revoke current refresh token (logout)"],
            ["POST", "/change-password", "Authorized", "Change password (requires current password)"],
            ["POST", "/forgot-password", "Public",     "Send password reset email"],
            ["POST", "/reset-password",  "Public",     "Complete reset with token + new password"],
        ],
        col_widths=[0.8, 2.0, 1.2, 2.6]
    )

    heading2(doc, "6.2  Items  (/api/v1/items)")
    add_table(doc,
        headers=["Method", "Path", "Auth", "Description"],
        rows=[
            ["GET",   "/",           "Public",      "Browse items (pagination, filters, sort). Cached 30s."],
            ["POST",  "/",           "Authorized",  "Create a new item listing"],
            ["GET",   "/{id}",       "Public",      "Get item details; increments view count"],
            ["PATCH", "/{id}",       "Owner",       "Update item fields"],
            ["DELETE","/{id}",       "Owner",       "Delete listing"],
            ["POST",  "/{id}/like",  "Authorized",  "Toggle like/unlike on an item"],
            ["POST",  "/{id}/bump",  "Owner",       "Re-surface item (re-set createdAt)"],
            ["POST",  "/ai-suggest", "Authorized",  "Submit photo → receive AI-generated listing suggestions"],
        ],
        col_widths=[0.8, 1.7, 1.2, 3.0]
    )

    heading2(doc, "6.3  Payments  (/api/v1/payments)")
    add_table(doc,
        headers=["Method", "Path", "Auth", "Description"],
        rows=[
            ["POST", "/checkout/{itemId}",   "Authorized", "Create Stripe Checkout session for card payment"],
            ["POST", "/on-spot/{itemId}",    "Authorized", "Record an on-spot (cash) payment"],
            ["POST", "/webhook",             "Public*",    "Stripe webhook handler (Stripe-Signature verified)"],
            ["GET",  "/",                    "Authorized", "List the authenticated user's payments"],
            ["POST", "/wallet/topup",        "Authorized", "Create PaymentIntent to top up internal wallet"],
        ],
        col_widths=[0.8, 2.2, 1.2, 2.5]
    )

    heading2(doc, "6.4  Shipping  (/api/v1/shipping)")
    add_table(doc,
        headers=["Method", "Path", "Auth", "Description"],
        rows=[
            ["POST", "/calculate",             "Authorized", "Calculate shipping price for chosen courier and delivery type"],
            ["POST", "/create",                "Authorized", "Create waybill with chosen courier"],
            ["GET",  "/offices/{provider}",    "Authorized", "List offices / lockers for a given courier"],
            ["GET",  "/{shipmentId}",          "Authorized", "Get shipment details and current status"],
            ["GET",  "/label/{shipmentId}",    "Authorized", "Download shipping label PDF"],
            ["POST", "/{shipmentId}/track",    "Authorized", "Sync status from courier API"],
            ["POST", "/{shipmentId}/cancel",   "Authorized", "Cancel an active shipment"],
        ],
        col_widths=[0.8, 2.2, 1.2, 2.5]
    )

    heading2(doc, "6.5  Additional Endpoints")
    add_table(doc,
        headers=["Controller", "Base Path", "Key Endpoints"],
        rows=[
            ["Messages",             "/api/v1/messages",               "GET conversations, GET /userId (history), POST /mark-read"],
            ["E-Bills",              "/api/v1/ebills",                 "GET list, POST /{id}/resend (rate-limited)"],
            ["Bundles",              "/api/v1/bundles",                "CRUD + purchase requests"],
            ["Users",                "/api/v1/users",                  "GET profile, PATCH self, GET /{id}/items, GET /{id}/ratings"],
            ["Categories",           "/api/v1/categories",             "GET list (cached 1h)"],
            ["Offers",               "/api/v1/offers",                 "POST create, PATCH counter/accept/decline"],
            ["Purchase Requests",    "/api/v1/purchase-requests",      "POST create, POST accept/decline"],
            ["Follows",              "/api/v1/follows",                "POST follow, DELETE unfollow, GET followers/following"],
            ["Ratings",              "/api/v1/ratings",                "POST rate user post-purchase"],
            ["Doctor Reviews",       "/api/v1/doctor-reviews",         "POST submit, GET approved list (admin: approve/reject)"],
            ["Child-Friendly Places","/api/v1/child-friendly-places",  "POST submit, GET filtered list (admin: approve/reject)"],
            ["Saved Searches",       "/api/v1/saved-searches",         "POST create, GET list, DELETE remove"],
            ["Feedback",             "/api/v1/feedback",               "POST submit, admin: GET list"],
            ["Photos",               "/api/v1/photos",                 "POST upload (R2), DELETE remove"],
            ["Admin",                "/api/v1/admin",                  "Dashboard stats, user block/unblock, item moderation, audit logs"],
            ["Assistant",            "/api/v1/assistant",              "POST chat message, GET conversation history"],
        ],
        col_widths=[1.9, 2.2, 2.5]
    )


def section_automation(doc):
    heading1(doc, "7. Automation & Workflows (n8n)")

    body(doc,
        "MamVibe ships with 16 pre-built n8n workflow files in the n8n-workflows/ directory. "
        "These are imported into a self-hosted n8n instance and triggered via HTTP webhooks "
        "dispatched by the platform's N8nWebhookService background service."
    )

    add_table(doc,
        headers=["Event", "Trigger", "Automated Action"],
        rows=[
            ["payment.completed",        "Stripe checkout confirmed",          "Send purchase confirmation emails to buyer and seller"],
            ["shipment.created",         "New waybill created",               "Email buyer tracking number and courier details"],
            ["shipment.delivered",       "Courier status → Delivered",        "Email both parties; prompt buyer for rating"],
            ["shipment.stuck",           "Daily: in transit 7 + days",        "Alert admin; notify buyer with tracking link"],
            ["user.registered",          "New account created",               "Welcome email with platform guide"],
            ["user.blocked",             "Admin blocks account",              "Notification email to affected user"],
            ["chat.message_offline",     "Message sent to offline user",      "Email excerpt of message with link to reply"],
            ["stale_items",             "Daily: listed 30 + days",           "Nudge seller to review, discount, or remove"],
            ["daily_summary",            "Daily at 08:00 UTC",               "Admin report: new items, payments, shipments"],
            ["feedback_prompt",          "Daily: delivered 2 + days, no rating", "Prompt buyer for seller rating"],
            ["weekly_seller_report",     "Monday 09:00 AM",                  "Email seller: sales summary, revenue, avg rating"],
        ],
        col_widths=[2.0, 2.0, 2.6]
    )


def section_security(doc):
    heading1(doc, "8. Security & Compliance")

    heading2(doc, "8.1  Security Controls")
    add_table(doc,
        headers=["Control", "Implementation"],
        rows=[
            ["Transport security",         "HTTPS enforced; HSTS header (max-age=31536000)"],
            ["API authentication",         "JWT Bearer tokens; short access token lifetime; HttpOnly refresh cookies"],
            ["OAuth",                      "Google OAuth 2.0 via ASP.NET Identity external login"],
            ["CAPTCHA",                    "Cloudflare Turnstile on registration, login, and password reset"],
            ["Password policy",            "Min 8 chars, uppercase, lowercase, digit, special character (ASP.NET Identity)"],
            ["Account lockout",            "5 failed logins → 5-minute lockout"],
            ["Rate limiting",              "Global 200 req/min; auth 30/min; upload 20/min; e-bill resend 3/min"],
            ["Input validation",           "FluentValidation on all request DTOs; no raw entity binding"],
            ["CORS",                       "Explicit allow-list of trusted origins; no wildcard"],
            ["HTTP security headers",      "CSP, X-Frame-Options (DENY), X-Content-Type-Options, Referrer-Policy, Permissions-Policy, CORP, COOP"],
            ["Blocked user enforcement",   "Middleware rejects all requests from blocked accounts on every call"],
            ["Metrics protection",         "/metrics endpoint returns 404 for non-internal IP addresses"],
            ["Stripe webhook validation",  "Stripe-Signature header verified before any payment processing"],
            ["Data protection keys",       "ASP.NET Data Protection keys persisted to Docker volume (survive restarts)"],
            ["Sensitive data in logs",     "Only entity IDs and status codes logged; no PII or financial data"],
            ["GDPR",                       "Data export endpoint lets users download all their personal data"],
            ["Swagger",                    "API documentation served only in Development environment"],
        ],
        col_widths=[2.2, 4.4]
    )

    heading2(doc, "8.2  Dependency Security")
    body(doc,
        "All NuGet packages are regularly audited using dotnet list package --vulnerable. "
        "No known vulnerable packages are in the dependency tree. Frontend npm packages "
        "are audited via npm audit."
    )


def section_localisation(doc):
    heading1(doc, "9. Localisation & Accessibility")

    heading2(doc, "9.1  Language Support")
    body(doc,
        "MamVibe supports English and Bulgarian across the entire frontend. Language detection "
        "follows this priority: user preference stored in their account → localStorage → "
        "browser language → fallback to English."
    )
    add_table(doc,
        headers=["Locale", "Coverage"],
        rows=[
            ["English (en)", "Full UI translation — all labels, messages, and error strings"],
            ["Bulgarian (bg)", "Full UI translation — all labels, messages, and error strings"],
        ],
        col_widths=[1.5, 5.1]
    )

    heading2(doc, "9.2  Fiscal & Legal Localisation")
    for f in [
        "Bulgarian VAT rate of 20 % applied to all taxable transactions.",
        "TakeANap fiscal receipt integration for Bulgarian tax authority compliance.",
        "E-bill number format follows local numbering conventions.",
        "Date and number formatting adapts to the selected locale via date-fns.",
        "Privacy Policy, Terms of Service, and Cookie Policy pages in both languages.",
    ]:
        bullet(doc, f)


def section_deployment(doc):
    heading1(doc, "10. Deployment & Operations")

    heading2(doc, "10.1  Docker Compose Stack")
    add_table(doc,
        headers=["Service", "Image / Tech", "Resources", "Role"],
        rows=[
            ["postgres",    "PostgreSQL 18",          "512 MB RAM, 1 CPU", "Primary database with health check"],
            ["pgbouncer",   "PgBouncer",              "Light",             "Connection pooler (transaction mode, 500 clients)"],
            ["redis",       "Redis 7",                "256 MB LRU + AOF",  "Cache + SignalR backplane"],
            ["api",         ".NET 8 / ASP.NET Core",  "1 GB RAM, 2 CPU",   "Backend REST API + SignalR hub"],
            ["frontend",    "Nginx + React build",   "Light",             "Static file serving + reverse proxy"],
            ["prometheus",  "Prometheus",             "256 MB",           "Metrics scraper"],
            ["loki",        "Grafana Loki",           "256 MB",           "Log aggregation (31-day retention)"],
            ["grafana",     "Grafana",                "256 MB",           "Monitoring dashboards and alerts"],
        ],
        col_widths=[1.3, 1.9, 1.5, 2.0]
    )

    heading2(doc, "10.2  Quick-start Commands")
    heading3(doc, "Full Stack (Docker)")
    for cmd in [
        "docker compose up --build    # Build and start all services",
        "docker compose down          # Stop all services",
    ]:
        bullet(doc, cmd)

    heading3(doc, "Development Mode (local services)")
    for cmd in [
        "docker compose up postgres -d              # Start database only",
        "dotnet run  (from backend/src/MomVibe.WebApi/)   # API at localhost:5038",
        "npm run dev (from frontend/)               # UI at localhost:5173",
    ]:
        bullet(doc, cmd)

    heading2(doc, "10.3  Environment Variables")
    body(doc,
        "All secrets are supplied via a root-level .env file (see .env.example for the full template). "
        "Doppler integration is available for production secret management."
    )
    add_table(doc,
        headers=["Variable Group", "Key Variables"],
        rows=[
            ["PostgreSQL",       "POSTGRES_DB, POSTGRES_USER, POSTGRES_PASSWORD, POSTGRES_PORT"],
            ["JWT",              "JWT_SECRET (minimum 32 characters)"],
            ["OAuth",            "GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET, FRONTEND_URL"],
            ["Stripe",           "STRIPE_SECRET_KEY, STRIPE_WEBHOOK_SECRET, STRIPE_WALLET_WEBHOOK_SECRET"],
            ["Couriers",         "ECONT_USERNAME/PASSWORD, SPEEDY_USERNAME/PASSWORD, BOXNOW_API_KEY, PIGEONEXPRESS_API_KEY"],
            ["Fiscal",           "TAKEANAP_API_KEY, TAKEANAP_API_SECRET, TAKEANAP_SHOP_ID"],
            ["Email / SMTP",     "SMTP_HOST, SMTP_PORT, SMTP_USERNAME, SMTP_PASSWORD, SMTP_FROM_EMAIL"],
            ["AI",               "ANTHROPIC_API_KEY, GROQ_API_KEY, AI_CHAT_PROVIDER (anthropic|groq)"],
            ["Cloudflare",       "TURNSTILE_SITE_KEY, TURNSTILE_SECRET_KEY, R2 credentials"],
            ["n8n",              "N8N_BASE_URL, N8N_ENABLED, N8N_WEBHOOK_SECRET"],
            ["Observability",    "OpenTelemetry__Otlp__Endpoint"],
        ],
        col_widths=[1.7, 4.9]
    )

    heading2(doc, "10.4  Branching Strategy (GitFlow)")
    add_table(doc,
        headers=["Branch", "Purpose"],
        rows=[
            ["main",          "Always deployable. Only production-ready, tested code is merged here."],
            ["develop",       "Integration branch. Feature branches merge here first."],
            ["feature/*",     "Branch from develop; one branch per feature or bug fix."],
            ["release/x.y.z", "Stabilisation before production. Merged to both main and develop."],
            ["hotfix/*",      "Emergency production fix. Branched from main; merged to main and develop."],
        ],
        col_widths=[1.6, 5.0]
    )


def section_team_guide(doc):
    heading1(doc, "11. Developer Onboarding Guide")

    heading2(doc, "11.1  Repository Layout")
    add_table(doc,
        headers=["Path", "Contents"],
        rows=[
            ["backend/src/MomVibe.Domain/",         "Entities, enums, domain constants"],
            ["backend/src/MomVibe.Application/",    "DTOs, interfaces, validators, AutoMapper profiles"],
            ["backend/src/MomVibe.Infrastructure/", "EF Core, migrations, external service clients"],
            ["backend/src/MomVibe.WebApi/",         "Controllers, SignalR hubs, middleware, StartUp.cs"],
            ["backend/tests/MomVibe.UnitTests/",    "xUnit unit tests (mocked dependencies)"],
            ["backend/tests/MomVibe.IntegrationTests/", "xUnit integration tests (real DB via WebApplicationFactory)"],
            ["frontend/src/",                       "React 19 app source"],
            ["n8n-workflows/",                      "16 pre-built n8n workflow JSON files"],
            ["docs/",                               "ADRs, screenshots, this document"],
        ],
        col_widths=[3.0, 3.6]
    )

    heading2(doc, "11.2  Common Commands")
    heading3(doc, "Backend")
    for cmd in [
        "dotnet run                                          # Start API (from MomVibe.WebApi/)",
        "dotnet test backend/tests/MomVibe.UnitTests        # Run unit tests",
        "dotnet test backend/tests/MomVibe.IntegrationTests # Run integration tests",
        "dotnet ef migrations add <Name> --project ../MomVibe.Infrastructure --startup-project .",
        "dotnet ef database update --project ../MomVibe.Infrastructure --startup-project .",
    ]:
        bullet(doc, cmd)

    heading3(doc, "Frontend")
    for cmd in [
        "npm run dev      # Start Vite dev server at localhost:5173",
        "npm run build    # TypeScript check + production bundle",
        "npm run lint     # ESLint",
        "npm run preview  # Preview production build locally",
    ]:
        bullet(doc, cmd)

    heading2(doc, "11.3  Naming Convention Note")
    body(doc,
        "The repository and Docker Compose file use the name MamVibe. "
        "However, all .NET project names, namespaces, and C# source code use MomVibe "
        "(e.g., MomVibe.Domain, MomVibe.Application). New code should follow the "
        "existing convention: namespace MomVibe.*"
    )

    heading2(doc, "11.4  Adding a New Feature — Checklist")
    for step in [
        "Create or update domain entities in MomVibe.Domain/.",
        "Add or update the interface in MomVibe.Application/Interfaces/.",
        "Add the request/response DTO and FluentValidation validator in MomVibe.Application/.",
        "Implement the service and any external API clients in MomVibe.Infrastructure/.",
        "Register services in the appropriate AddInfrastructureServices() extension method.",
        "Add the controller endpoint in MomVibe.WebApi/Controllers/.",
        "Write unit tests in MomVibe.UnitTests/ and integration tests in MomVibe.IntegrationTests/.",
        "Add or update frontend API client in frontend/src/api/.",
        "Update Zustand store if global state is needed.",
        "Add the page/component in frontend/src/pages/ or frontend/src/components/.",
        "Add translation keys to both en.json and bg.json.",
        "Add a new EF Core migration if the schema changed.",
    ]:
        bullet(doc, step)


def section_roadmap(doc):
    heading1(doc, "12. Current Status & Roadmap")

    heading2(doc, "12.1  Completed (Production-Ready)")
    for item in [
        "Full marketplace: create, browse, buy, and sell listings.",
        "Donate and sell listing types with AI photo assistant.",
        "AI content moderation with confidence scores.",
        "Complete payment stack: Stripe, Wallet, On-spot, COD.",
        "Full shipping integration: Econt, Speedy, Box Now, Pigeon Express.",
        "Real-time SignalR chat with typing indicators and read receipts.",
        "AI chat bot (Anthropic Claude / Groq).",
        "E-bills and TakeANap fiscal receipts.",
        "Admin panel: users, items, shipments, audit logs.",
        "Doctor reviews and child-friendly places community directories.",
        "Follow/follower system and post-purchase ratings.",
        "Saved searches with email alerts.",
        "n8n automation: 16 pre-built workflows.",
        "Bilingual UI (English and Bulgarian).",
        "Full observability stack (Prometheus, Loki, Grafana).",
        "Cloudflare Turnstile bot protection.",
        "GDPR data export.",
        "Docker Compose production deployment.",
    ]:
        bullet(doc, item)

    heading2(doc, "12.2  Potential Next Steps")
    for item in [
        "Native mobile application (React Native / Expo) with full push notification support.",
        "Expanded courier integrations for additional EU markets.",
        "Seller subscription tiers (featured listings, priority support).",
        "Advanced search: map-based item browse, radius filtering.",
        "Gamification: seller milestones, badges, loyalty points.",
        "Expanded AI features: automated price benchmarking, listing quality scores.",
        "Multi-country localisation (Greek, Romanian) for Balkans expansion.",
    ]:
        bullet(doc, item)


def section_glossary(doc):
    heading1(doc, "13. Glossary")

    add_table(doc,
        headers=["Term", "Definition"],
        rows=[
            ["MamVibe",         "The brand and repository name for the platform."],
            ["MomVibe",         "The namespace used in all .NET code (MomVibe.Domain, etc.)"],
            ["Item",            "A single product listed for donation or sale."],
            ["Bundle",          "A group of 2-10 items sold together at a combined discounted price."],
            ["Listing Type",    "Donate (free) or Sell (priced)."],
            ["Purchase Request","A formal buyer request to acquire an item; seller must accept."],
            ["Offer",           "A price proposal below the listed price; seller can counter."],
            ["Wallet",          "Internal escrow account holding buyer funds until delivery is confirmed."],
            ["E-Bill",          "Electronic payment receipt auto-generated for every completed purchase."],
            ["Waybill",         "Courier shipping document; unique tracking reference for a parcel."],
            ["TakeANap",        "Bulgarian fiscal receipt provider integrated for VAT compliance."],
            ["n8n",             "Open-source workflow automation tool used for email and alert workflows."],
            ["Turnstile",       "Cloudflare's invisible CAPTCHA service used on authentication forms."],
            ["PgBouncer",       "PostgreSQL connection pooler running in transaction mode."],
            ["SignalR",         "Microsoft's WebSocket library used for real-time chat and notifications."],
            ["R2",              "Cloudflare's S3-compatible object storage used for photo uploads."],
            ["Groq",            "Low-latency AI inference provider used as an alternative to Anthropic."],
            ["GitFlow",         "Branching strategy: main, develop, feature/*, release/*, hotfix/*."],
            ["Clean Architecture","Layered code structure: Domain → Application → Infrastructure → WebApi."],
            ["COD",             "Cash on Delivery — courier collects payment at the recipient's door."],
        ],
        col_widths=[1.8, 4.8]
    )


# ──────────────────────────────────────────────────────────
# Main builder
# ──────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # Default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = WARM_GRAY

    # ── Cover ──────────────────────────────────────────────
    build_cover(doc)

    # ── Table of Contents placeholder ─────────────────────
    heading1(doc, "Table of Contents")
    toc_items = [
        ("1.",  "Executive Summary"),
        ("2.",  "Product Overview"),
        ("3.",  "Feature Set"),
        ("4.",  "Technology Stack"),
        ("5.",  "Architecture"),
        ("6.",  "API Reference Summary"),
        ("7.",  "Automation & Workflows (n8n)"),
        ("8.",  "Security & Compliance"),
        ("9.",  "Localisation & Accessibility"),
        ("10.", "Deployment & Operations"),
        ("11.", "Developer Onboarding Guide"),
        ("12.", "Current Status & Roadmap"),
        ("13.", "Glossary"),
    ]
    for num, title in toc_items:
        p    = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        rnum = p.add_run(f"{num}  ")
        rnum.bold = True
        rnum.font.size = Pt(10.5)
        rnum.font.color.rgb = PURPLE
        rtitle = p.add_run(title)
        rtitle.font.size = Pt(10.5)
        rtitle.font.color.rgb = WARM_GRAY

    page_break(doc)

    # ── Sections ───────────────────────────────────────────
    section_executive_summary(doc)
    page_break(doc)

    section_product_overview(doc)
    page_break(doc)

    section_features(doc)
    page_break(doc)

    section_tech_stack(doc)
    page_break(doc)

    section_architecture(doc)
    page_break(doc)

    section_api_reference(doc)
    page_break(doc)

    section_automation(doc)
    section_security(doc)
    page_break(doc)

    section_localisation(doc)
    section_deployment(doc)
    page_break(doc)

    section_team_guide(doc)
    page_break(doc)

    section_roadmap(doc)
    section_glossary(doc)

    # ── Footer note ────────────────────────────────────────
    page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 60)
    run.font.color.rgb = PURPLE_LIGHT
    run.font.size = Pt(10)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(
        f"MamVibe Platform Overview  ·  Confidential  ·  {datetime.date.today().year}"
    )
    r2.font.size = Pt(9)
    r2.font.color.rgb = WARM_GRAY
    r2.italic = True

    return doc


# ──────────────────────────────────────────────────────────
# Entry point
# ──────────────────────────────────────────────────────────

if __name__ == "__main__":
    output_path = r"C:\WORK_PLACE\MamVibe\docs\word_document\MamVibe_Platform_Overview.docx"
    doc = build_document()
    doc.save(output_path)
    print(f"Document saved: {output_path}")
