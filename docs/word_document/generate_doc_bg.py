"""
MamVibe — Преглед на платформата (Български)
Генерира изчерпателен .docx отчет за платформата MamVibe на български език.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ──────────────────────────────────────────────────────────
# Цветова палитра
# ──────────────────────────────────────────────────────────
PURPLE       = RGBColor(0x7C, 0x3A, 0xED)
PURPLE_DARK  = RGBColor(0x5B, 0x21, 0xB6)
PURPLE_LIGHT = RGBColor(0xED, 0xE9, 0xFE)
WARM_GRAY    = RGBColor(0x6B, 0x72, 0x80)
ACCENT_PINK  = RGBColor(0xF4, 0x72, 0xB6)
ACCENT_TEAL  = RGBColor(0x14, 0xB8, 0xA6)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
DARK         = RGBColor(0x1F, 0x2A, 0x3C)
LIGHT_GRAY   = RGBColor(0xF9, 0xFA, 0xFB)

# ──────────────────────────────────────────────────────────
# XML помощни функции
# ──────────────────────────────────────────────────────────

def rgb_to_hex(rgb: RGBColor) -> str:
    return str(rgb)


def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  rgb_to_hex(rgb))
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),   kwargs.get("val",   "single"))
        tag.set(qn("w:sz"),    kwargs.get("sz",    "4"))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color", "E5E7EB"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def set_run_font(run, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def page_break(doc):
    doc.add_page_break()


def add_horizontal_rule(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "7C3AED")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(12)

# ──────────────────────────────────────────────────────────
# Стилизирани параграфи
# ──────────────────────────────────────────────────────────

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = PURPLE_DARK
    set_run_font(run)
    add_horizontal_rule(doc)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = PURPLE
    set_run_font(run)
    return p


def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK
    set_run_font(run)
    return p


def body(doc, text, italic=False, color=None):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.italic    = italic
    run.font.color.rgb = color or WARM_GRAY
    set_run_font(run)
    return p


def bullet(doc, text, level=0):
    p   = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = WARM_GRAY
    set_run_font(run)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(3)
    return p


def add_table(doc, headers, rows, col_widths=None, alt_rows=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style     = "Table Grid"

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, PURPLE_DARK)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p    = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        set_run_font(run)

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg  = PURPLE_LIGHT if (alt_rows and r_idx % 2 == 0) else WHITE
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p    = cell.paragraphs[0]
            run  = p.add_run(cell_text)
            run.font.size = Pt(10)
            run.font.color.rgb = WARM_GRAY
            set_run_font(run)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    doc.add_paragraph()
    return table

# ──────────────────────────────────────────────────────────
# Корица
# ──────────────────────────────────────────────────────────

def build_cover(doc):
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, PURPLE_DARK)
    cell.width = Inches(6.5)
    cell.paragraphs[0].clear()

    def cover_para(text, size, color=WHITE, bold=False, space_before=0, space_after=6, center=True):
        p   = cell.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        run = p.add_run(text)
        run.bold = bold
        run.font.size  = Pt(size)
        run.font.color.rgb = color
        set_run_font(run)
        return p

    cover_para("", 8, space_before=60, space_after=0)
    cover_para("✦  MamVibe  ✦", 13, color=ACCENT_PINK, bold=False, space_after=4)
    cover_para("Преглед на платформата", 38, bold=True, space_before=8, space_after=8)
    cover_para("─" * 45, 10, color=PURPLE_LIGHT, space_after=16)
    cover_para(
        "Модерният пазар, където семействата се свързват,\n"
        "даряват и пазаруват детски стоки и аксесоари.",
        13, color=PURPLE_LIGHT, space_after=32
    )
    cover_para(f"Дата на доклада:  {datetime.date.today().strftime('%d %B %Y')}", 10, color=PURPLE_LIGHT, space_after=4)
    cover_para("Версия:  1.0", 10, color=PURPLE_LIGHT, space_after=4)
    cover_para("Поверително — За вътрешно и инвеститорско ползване", 9, color=ACCENT_PINK, space_after=60)

    doc.add_paragraph()
    page_break(doc)

# ──────────────────────────────────────────────────────────
# Раздели
# ──────────────────────────────────────────────────────────

def section_executive_summary(doc):
    heading1(doc, "1. Резюме")

    body(doc,
        "MamVibe е цялостна дигитална платформа — пазар, предназначена за семейства с малки деца. "
        "Тя предоставя сигурна, общностно ориентирана среда, в която родителите могат да даряват, "
        "купуват или продават употребявани бебешки и детски вещи — от дрехи и колички до играчки "
        "и аксесоари. Платформата съчетава богато електронно-търговско преживяване с чат в реално "
        "време, инструменти с изкуствен интелект, интегрирана логистика и жива общност."
    )

    body(doc,
        "Изградена с технологии от корпоративен клас (.NET 8 за бекенда, React 19 за фронтенда), "
        "MamVibe е готова за продукция и проектирана да мащабира. Тя се интегрира с водещите "
        "български куриерски компании, Stripe за картови плащания, вътрешна ескроу-портфейл система, "
        "Cloudflare за сигурност и съхранение, и AI доставчици (Anthropic Claude, Groq) за умен "
        "асистент при публикуване и модерация на съдържание."
    )

    body(doc,
        "MamVibe е насочена основно към български семейства с двуезична поддръжка "
        "(английски и български), съответствие с българското фискално законодателство "
        "(интеграция с TakeANap, ДДС 20%) и интеграции с четирите най-популярни местни куриери. "
        "Архитектурата позволява разширяване на нови пазари с минимални инженерни усилия."
    )

    heading2(doc, "С един поглед")
    add_table(doc,
        headers=["Измерение", "Детайл"],
        rows=[
            ["Вид платформа",       "Пазар между физически лица (C2C) + общностен хъб"],
            ["Основен пазар",       "България (двуезично: EN / BG)"],
            ["Целева аудитория",    "Семейства с деца от 0 до 16 години"],
            ["Видове обяви",        "Дарение (безплатно) и Продажба (с цена)"],
            ["Методи за плащане",   "Карта (Stripe), Портфейл ескроу, На място, Наложен платеж"],
            ["Куриерски компании",  "Еконт, Speedy, Box Now, Pigeon Express"],
            ["Функции в реално вр.","SignalR чат, индикатори за писане, пуш известия"],
            ["AI функции",          "Предложения за обяви, модерация на съдържание, AI чат бот"],
            ["Бекенд",              ".NET 8 / ASP.NET Core — Чиста архитектура"],
            ["Фронтенд",            "React 19 + TypeScript + Vite + Tailwind CSS"],
            ["База данни",          "PostgreSQL 18 + PgBouncer + Redis"],
            ["Инфраструктура",      "Docker Compose, Nginx, Grafana, Prometheus, Loki"],
        ],
        col_widths=[2.2, 4.3]
    )


def section_product_overview(doc):
    heading1(doc, "2. Преглед на продукта")

    body(doc,
        "MamVibe е създадена с убеждението, че детските вещи — използвани само за няколко месеца — "
        "заслужават втори живот. Вместо да завършат в кофите за боклук, колички, дрехи и играчки "
        "могат да намерят ново място в доверена родителска общност. MamVibe прави тази размяна "
        "лесна, сигурна и социално значима."
    )

    heading2(doc, "2.1  Мисия")
    body(doc,
        "Да изградим най-доверената и удобна платформа за обмен на детски вещи между семействата, "
        "насърчавайки устойчивостта, общността и финансовите спестявания за родителите в България."
    )

    heading2(doc, "2.2  Основни потребителски пътища")

    heading3(doc, "Пътят на продавача / дарителя")
    for step in [
        "Регистрирайте се и настройте профил (индивидуален, двойка или семейство).",
        "Създайте обява: качете до 5 снимки, задайте заглавие, описание, категория, състояние, възрастова група, размер на дрехи/обувки и цена (или маркирайте като безплатно дарение).",
        "AI асистентът анализира снимката и попълва предварително данните на обявата, за да спести време.",
        "Обявата излиза на живо след преминаване на AI модерация и (по желание) преглед от администратор.",
        "Получавайте заявки за покупка или директно плащане чрез Stripe от купувачи.",
        "Приемете заявката → пратката се създава автоматично с избрания куриер.",
        "Средствата се задържат в ескроу и се освобождават в портфейла на продавача след потвърдена доставка.",
        "Продавачът получава оценка от купувача.",
    ]:
        bullet(doc, step)

    heading3(doc, "Пътят на купувача")
    for step in [
        "Разгледайте обяви с филтри: категория, вид обява, възрастова група, размер, състояние, ценови диапазон.",
        "Харесвайте и запазвайте обяви; създавайте запазени търсения с имейл известия за нови съвпадения.",
        "Изпратете заявка за покупка или купете незабавно чрез Stripe карта.",
        "Изберете доставка: офис на куриер, домашен адрес или автомат.",
        "Проследявайте пратката в реално време от таблото за управление.",
        "Потвърдете получаването → средствата се освобождават на продавача.",
        "Оставете звездна оценка и коментар за продавача.",
    ]:
        bullet(doc, step)

    heading3(doc, "Общностният път")
    for step in [
        "Следвайте други потребители и разглеждайте техните обяви.",
        "Общувайте в реално време с продавачи или купувачи.",
        "Откривайте и рецензирайте местни лекари и детски места (детски площадки, кафенета, паркове).",
        "Подавайте обратна връзка за платформата (докладвайте грешки, предлагайте функции).",
        "Достъп до двуезично съдържание (EN / BG) според личните предпочитания.",
    ]:
        bullet(doc, step)


def section_features(doc):
    heading1(doc, "3. Функционалности")

    heading2(doc, "3.1  Основни функции на пазара")
    for f in [
        "Обяви с до 5 снимки, заглавие, описание, категория, състояние, възрастова група, размери на дрехи и обувки.",
        "Два вида обяви: Дарение (безплатно) и Продажба (с цена).",
        "Разширено търсене и разглеждане: филтриране по категория, вид обява, възрастова група, размер, състояние, ценови диапазон и ключова дума.",
        "Страниране и множество режими на сортиране (най-нови, цена, популярност).",
        "Броячи на прегледи и харесвания/любими за всяка обява.",
        "Функция 'Освежи' (Bump): публикувайте отново обявата, за да я изведете на преден план (след 7 дни).",
        "Създаване на пакети (Bundle): групирайте от 2 до 10 артикула на комбинирана намалена цена.",
        "Работен процес за заявки за покупка: купувачът изпраща заявка → продавачът приема или отхвърля.",
        "Статус (Резервиран): маркирайте артикул като резервиран, без да го продавате напълно.",
        "Режим (Ваканция): временно скрийте всички обяви на продавача.",
        "Детекция на стари обяви: маркиране на обяви, активни повече от 30 дни.",
        "AI асистент за обяви: анализ на снимка → автоматично попълване на заглавие, описание, предложение за цена и категория.",
        "AI модерация на съдържание: автоматичен преглед на описанията с оценки на увереност и резервен преглед от администратор.",
    ]:
        bullet(doc, f)

    heading2(doc, "3.2  Плащания")
    add_table(doc,
        headers=["Метод на плащане", "Доставчик", "Процес"],
        rows=[
            ["Карта",            "Stripe",          "Stripe Checkout сесия → потвърждение чрез webhook → средства в ескроу портфейл"],
            ["Портфейл",         "Вътрешен ескроу", "Купувачът зарежда портфейла си чрез Stripe PaymentIntent → средствата се задържат → освобождават се при доставка"],
            ["На място",         "Вътрешен",        "Продавачът записва ръчно транзакцията в брой; платформата следи без движение на средства"],
            ["Наложен платеж",   "Куриерски COD",   "Куриерът събира парите при доставка; платформата синхронизира статуса при потвърждение"],
            ["Резервация",       "Вътрешен",        "Безплатна резервация на дарение; не е необходима финансова транзакция"],
        ],
        col_widths=[1.8, 1.8, 3.0]
    )

    heading2(doc, "3.3  Доставка и логистика")
    body(doc,
        "MamVibe се интегрира с четирите основни български куриерски мрежи. Архитектурата на "
        "модулни доставчици позволява добавянето на нови куриери с минимални инженерни усилия."
    )
    add_table(doc,
        headers=["Куриер", "Видове доставка", "Функционалности"],
        rows=[
            ["Еконт Експрес",   "Офис / Адрес / Автомат", "Пълна API: изчисляване на цена, товарителница, проследяване, COD, застраховка"],
            ["Speedy",          "Офис / Адрес",            "Пълна API: изчисляване на цена, товарителница, проследяване, COD"],
            ["Box Now",         "Автомат / Офис",          "Мрежа от автомати, изчисляване на цена, проследяване"],
            ["Pigeon Express",  "Адрес / Офис",            "API интеграция, проследяване, COD"],
        ],
        col_widths=[1.6, 2.0, 3.0]
    )
    body(doc, "Общи функции за доставка:")
    for f in [
        "Изчисляване на цена на доставка по куриер, вид доставка и тегло на пратката.",
        "Автоматично генериране на товарителница и етикет; изтегляне на PDF от таблото за управление.",
        "Синхронизиране на проследяването в реално време: Изчакване → Създадена → Взета → В транзит → За доставка → Доставена.",
        "Опции за наложен платеж (COD) и застраховка.",
        "API за намиране на офиси и автомати за безадресно получаване.",
        "Административен преглед на пратките с управление на статусите в голям мащаб.",
    ]:
        bullet(doc, f)

    heading2(doc, "3.4  Съобщения и известия в реално време")
    for f in [
        "WebSocket чат, захранван от SignalR на адрес /hubs/chat.",
        "Индикатори за писане и потвърждения за прочитане.",
        "Проследяване на присъствие на потребители — значки онлайн/офлайн в целия интерфейс.",
        "Значки с брой непрочетени съобщения в навигационната лента.",
        "Имейл известие при офлайн съобщение: ако получателят не е свързан, се изпраща имейл.",
        "AI чат бот: специален системен потребител, захранван от Anthropic Claude или Groq, отговарящ на въпроси за обяви и платформата.",
        "Push известия от сървъра за заявки за покупка, актуализации на оферти, нови последователи и промени в статуса на пратките.",
        "Expo push известия за мобилни клиенти.",
    ]:
        bullet(doc, f)

    heading2(doc, "3.5  Общност и социални функции")
    for f in [
        "Потребителски профили: показвано ime, аватар (снимка или илюстрация), биография, тип профил (Мъж / Жена / Семейство).",
        "Система за следване/последователи с интеграция на социален фийд.",
        "Оценки след покупка (1-5 звезди) с писмени коментари; показват се в публичните профили.",
        "Рецензии за лекари: колективни отзиви за педиатри, зъболекари и специалисти — модерирани от администратор преди публикуване.",
        "Детски места: общностна директория на паркове, детски площадки, семейни ресторанти и др. — с филтри по възраст.",
        "Система за обратна връзка: потребителите подават доклади за грешки, предложения за функции или общи коментари със звездна оценка.",
    ]:
        bullet(doc, f)

    heading2(doc, "3.6  Финансово съответствие")
    for f in [
        "Електронни фактури (е-фактури) се генерират автоматично за всяко завършено плащане.",
        "Уникална и идемпотентна номерация на е-фактурите: MV-ГГГГ-{първи 8 hex символа на ID на плащането}.",
        "Доставка на е-фактури по имейл с ограничаване на скоростта (3 повторни изпращания/минута на потребител).",
        "Интеграция с TakeANap за фискален бон: HMAC-SHA256 подписани бонове, съответстващи на изискванията на НАП.",
        "Изчисляване на ДДС по стандартната българска ставка от 20%.",
    ]:
        bullet(doc, f)

    heading2(doc, "3.7  Административен панел")
    for f in [
        "Табло с живи статистики: активни обяви, чакащи плащания, пратки в транзит, последна обратна връзка.",
        "Управление на потребители: преглед на профили, блокиране/разблокиране на акаунти.",
        "Модерация на обяви: одобряване или отхвърляне; преглед на AI оценките за модерация.",
        "Надзор на доставките: пълен списък с пратки с филтри по статус и разбивки по куриери.",
        "Модерация на общността: одобряване на рецензии за лекари и детски места.",
        "Преглед на одитния журнал: филтруем запис на всички административни действия с времеви марки и разлики в промените.",
    ]:
        bullet(doc, f)

    heading2(doc, "3.8  Оферти и договаряне")
    for f in [
        "Купувачите могат да изпратят ценова оферта под обявената цена.",
        "Продавачите могат да направят контраоферта с различна цена.",
        "Жизнен цикъл на статуса на офертата: Чакаща → Приета / Отхвърлена.",
        "Приетите оферти автоматично се превръщат в заявка за покупка.",
    ]:
        bullet(doc, f)

    heading2(doc, "3.9  Запазени търсения и известия")
    for f in [
        "Запазете всяка конфигурация за търсене като именувано известие.",
        "Филтри: категория, вид обява, ключова дума, възрастова група, размер, състояние, максимална цена.",
        "Автоматичен имейл при нова обява, съответстваща на запазено търсене.",
        "Управление на запазените търсения от таблото на потребителя.",
    ]:
        bullet(doc, f)

    heading2(doc, "3.10  Сигурност и съответствие")
    for f in [
        "Cloudflare Turnstile CAPTCHA на всички форми за автентикация (регистрация, вход, възстановяване на парола).",
        "JWT токени за достъп (краткотрайни) + HttpOnly бисквитки за опресняващ токен.",
        "Единично влизане с Google OAuth.",
        "Политика за пароли на ASP.NET Identity: минимум 8 знака, главна буква, малка буква, цифра и специален знак.",
        "Блокиране при грешни опити: 5 неуспешни влизания водят до 5-минутно блокиране.",
        "Глобално ограничаване на заявките (200/мин), по-строги лимити за автентикация (30/мин) и качване (20/мин).",
        "GDPR: endpoint за експорт на данни — потребителите могат да изтеглят всичките си лични данни.",
        "Middleware за блокирани потребители: спрените акаунти се отхвърлят при всяка заявка.",
        "Пълен набор HTTP заглавки за сигурност: CSP, X-Frame-Options, HSTS, Referrer-Policy, Permissions-Policy, CORP, COOP.",
    ]:
        bullet(doc, f)


def section_tech_stack(doc):
    heading1(doc, "4. Технологичен стек")

    heading2(doc, "4.1  Бекенд")
    add_table(doc,
        headers=["Компонент", "Технология", "Предназначение"],
        rows=[
            ["Фреймуърк",          ".NET 8 / ASP.NET Core",        "Web API хост, dependency injection, middleware"],
            ["Архитектура",        "Чиста архитектура (4 слоя)",   "Domain → Application → Infrastructure → WebApi"],
            ["База данни",         "PostgreSQL 18",                 "Основно релационно хранилище за данни"],
            ["ORM",                "Entity Framework Core 9",       "Code-first миграции, LINQ заявки"],
            ["Connection pool",    "PgBouncer (transaction mode)",  "500 макс. клиента, 20 pool size — намалява DB връзките"],
            ["Кеш / PubSub",       "Redis 7",                       "Разпределен кеш, SignalR backplane"],
            ["В реално време",     "ASP.NET Core SignalR",          "WebSocket чат и push известия от сървъра"],
            ["Автентикация",       "ASP.NET Identity + JWT",        "Базирана на токени автентикация с опресняващи токени"],
            ["OAuth",              "Google OAuth 2.0",              "Единично влизане чрез Google акаунти"],
            ["Валидация",          "FluentValidation",              "Валидация на DTO заявки в слоя Application"],
            ["Mapping",            "AutoMapper",                    "Профили за映射 Entity ↔ DTO"],
            ["Логиране",           "Serilog",                       "Структурирано логиране с изход към Loki"],
            ["Наблюдаемост",       "OpenTelemetry",                 "Разпределени трасировки (OTLP export), метрики"],
            ["Фонови задачи",      "IHostedService / Channel<T>",   "n8n диспечер, синхр. на проследяване, планирани известия"],
            ["Версиониране на API","Asp.Versioning",                "Версиониране по URL (/api/v1/)"],
        ],
        col_widths=[1.8, 2.2, 2.6]
    )

    heading2(doc, "4.2  Фронтенд")
    add_table(doc,
        headers=["Компонент", "Технология", "Предназначение"],
        rows=[
            ["Фреймуърк",          "React 19 + TypeScript",      "Компонентен UI с пълна типова безопасност"],
            ["Инструмент за build","Vite",                       "Бърз HMR dev сървър, оптимизиран продукционен bundle"],
            ["Стилизиране",        "Tailwind CSS",               "Utility-first адаптивен дизайн"],
            ["Управление на стат.","Zustand",                    "Лека глобална среда с localStorage персистентност"],
            ["HTTP клиент",        "Axios",                      "API заявки с интерцептори (автентикация, обновяване, грешки)"],
            ["В реално време",     "@microsoft/signalr",         "WebSocket връзка към /hubs/chat"],
            ["Маршрутизиране",     "React Router v6",            "Клиентско маршрутизиране с защитени routes"],
            ["i18n",               "react-i18next",              "Преводи на потребителски интерфейс на EN и BG"],
            ["Известия",           "react-hot-toast",            "In-app toast известия"],
            ["Работа с дати",      "date-fns",                   "Форматиране на дати, съобразено с locale"],
            ["Path alias",         "@  →  src/",                 "Конфигуриран в vite.config.ts за чисти imports"],
            ["Dev proxy",          "Vite proxy",                 "Проксира /api, /hubs, /uploads към localhost:5038"],
        ],
        col_widths=[1.8, 2.2, 2.6]
    )

    heading2(doc, "4.3  Инфраструктура")
    add_table(doc,
        headers=["Услуга", "Технология", "Роля"],
        rows=[
            ["Контейнерна оркестр.", "Docker + Docker Compose",    "Разгръщане на целия стек локално и в продукция"],
            ["Обратен прокси",       "Nginx",                      "Обслужване на статични файлове, API проксиране, SSL"],
            ["Метрики",              "Prometheus",                 "Събира /metrics от API; съхранява времеви редове"],
            ["Логове",               "Loki (31-дневно задържане)", "Агрегира структурирани логове от всички услуги"],
            ["Табла и известия",     "Grafana",                    "Визуализира метрики + логове; изпраща предупреждения"],
            ["Съхранение на снимки", "Cloudflare R2 (S3-compat.)", "Object storage за снимки на обяви и аватари"],
            ["Защита от ботове",     "Cloudflare Turnstile",       "Невидим CAPTCHA при автентикация"],
            ["Управление на тайни",  "Doppler (по желание)",       "Централизирано инжектиране на тайни при разгръщане"],
            ["Data Protection ключ.","Docker volume",              "Персистира ASP.NET Data Protection ключовете"],
        ],
        col_widths=[2.0, 2.2, 2.4]
    )

    heading2(doc, "4.4  Интеграции с трети страни")
    add_table(doc,
        headers=["Интеграция", "Предназначение"],
        rows=[
            ["Stripe",               "Checkout сесии, PaymentIntent (зареждане на портфейл), webhooks"],
            ["Еконт Експрес API",    "Изчисляване на цена, товарителница, проследяване, COD"],
            ["Speedy API",           "Изчисляване на цена, товарителница, проследяване, COD"],
            ["Box Now API",          "Доставка до автомати, проследяване"],
            ["Pigeon Express API",   "Куриерска доставка, проследяване, COD"],
            ["TakeANap",             "Генериране на фискален бон за България (HMAC-SHA256 подписан)"],
            ["Anthropic Claude",     "AI асистент за обяви, модерация на съдържание, чат бот (vision + text)"],
            ["Groq",                 "Алтернативен AI доставчик за чат (ниска латентност)"],
            ["n8n",                  "Автоматизация: имейл известия, отчети за продавачи, предупреждения"],
            ["SMTP",                 "Транзакционен имейл: забравена парола, е-фактури, офлайн известия"],
            ["Google OAuth",         "Единично влизане за потребители без имейл/парола"],
            ["Cloudflare R2",        "S3-съвместимо object storage за качени от потребителите снимки"],
            ["Expo Push",            "Push известия до мобилни клиенти"],
        ],
        col_widths=[2.0, 4.5]
    )


def section_architecture(doc):
    heading1(doc, "5. Архитектура")

    heading2(doc, "5.1  Бекенд — Чиста архитектура")
    body(doc,
        "Бекендът е структуриран като четири отделни .NET проекта със строго наложен поток "
        "на зависимости: външните слоеве могат да зависят от вътрешните, но не и обратно."
    )
    add_table(doc,
        headers=["Слой", "Проект", "Съдържа"],
        rows=[
            ["Domain",         "MomVibe.Domain",         "Entities, enums, константи. Нула външни зависимости."],
            ["Application",    "MomVibe.Application",    "DTO, интерфейси на услуги, FluentValidation валидатори, AutoMapper профили."],
            ["Infrastructure", "MomVibe.Infrastructure", "EF Core DbContext + миграции, клиенти на външни API (Stripe, куриери, AI), имплементации на услуги."],
            ["Presentation",   "MomVibe.WebApi",         "Контролери, SignalR хъбове, middleware pipeline, DI окабеляване (StartUp.cs)."],
        ],
        col_widths=[1.4, 2.0, 3.2]
    )

    heading2(doc, "5.2  Ключови архитектурни шаблони")

    heading3(doc, "N8n Webhook Диспечер (Fire-and-forget)")
    body(doc,
        "Бизнес събитията (завършено плащане, доставена пратка, регистриран потребител и др.) "
        "се поставят в опашка чрез ограничен Channel<T> с капацитет 500. BackgroundService "
        "изчерпва опашката асинхронно и изпраща POST заявки с данните до конфигурирания n8n "
        "webhook URL. Това отделя пътя на заявката от латентността на външната автоматизация."
    )

    heading3(doc, "Фабрика за куриерски доставчици")
    body(doc,
        "CourierProviderFactory разрешава правилната имплементация на ICourierProvider "
        "(Еконт, Speedy, Box Now, Pigeon Express) въз основа на enum стойността в entity-то "
        "Shipment. Добавянето на нов куриер изисква само нов клас, имплементиращ интерфейса, "
        "и регистрация в DI контейнера — без промени в съществуващия код."
    )

    heading3(doc, "Идемпотентност при плащания")
    body(doc,
        "Незадължителното HTTP заглавие Idempotency-Key се съхранява в entity-то Payment. "
        "Уникален индекс в базата данни предотвратява дублирани редове Payment. Ключът се "
        "препраща към Stripe, за да предотврати дублирани такси дори при повторни опити."
    )

    heading3(doc, "Output Cache за endpoints с висок трафик")
    body(doc,
        "Endpoint-ът за разглеждане на обяви се кешира за 30 секунди чрез ASP.NET Output Cache. "
        "Категориите се кешират за 1 час с инвалидиране на базата на тагове. Redis е "
        "поддържащото хранилище, осигуряващо последователност при множество API инстанции."
    )

    heading3(doc, "SignalR с Redis Backplane")
    body(doc,
        "Всички събития в реално време преминават през SignalR. Redis backplane гарантира, "
        "че съобщенията и известията се разпространяват правилно дори когато работят "
        "множество API контейнерни инстанции зад Nginx балансьора на натоварване."
    )

    heading3(doc, "Идемпотентност на е-фактурите")
    body(doc,
        "Номерата на е-фактурите се присвояват точно веднъж. Колоната EBillNumber се "
        "проверява преди присвояване: ако вече е зададена (напр. от повторно извикване на "
        "Stripe webhook), втори номер не се генерира. Формат: MV-{ГОДИНА}-{първи 8 hex символа на payment ID}."
    )

    heading2(doc, "5.3  Схема на базата данни — ключови entities")
    add_table(doc,
        headers=["Entity", "Ключови полета", "Бележки"],
        rows=[
            ["ApplicationUser",     "DisplayName, ProfileType, AvatarUrl, IsBlocked, Bio, LanguagePreference, RevolutTag, ExpoPushToken, IsOnHoliday", "Разширява ASP.NET Identity"],
            ["Item",                "Title, Description, CategoryId, ListingType, AgeGroup, Size, Price, UserId, IsActive, IsReserved, IsSold, ViewCount, Condition, AiModerationStatus", "Основен marketplace entity"],
            ["Payment",             "ItemId, BuyerId, SellerId, Amount, PaymentMethod, PaymentStatus, StripeSessionId, EBillNumber, IdempotencyKey", "Покрива всички методи на плащане"],
            ["Shipment",            "PaymentId, CourierProvider, DeliveryType, Status, TrackingNumber, WaybillId, ShippingPrice, IsCod, CodAmount, IsInsured", "Свързан 1:1 с Payment"],
            ["Message",             "SenderId, ReceiverId, Content, IsRead", "Чат съобщение; прочита се в SignalR хъба"],
            ["PurchaseRequest",     "ItemId, BuyerId, SellerId, Status", "Чакаща / Приета / Отхвърлена"],
            ["Offer",               "ItemId, BuyerId, SellerId, OfferedPrice, CounterPrice, Status", "Entity за договаряне на цена"],
            ["UserRating",          "RaterId, RatedUserId, PurchaseRequestId, Rating, Comment", "Отзив след транзакция"],
            ["DoctorReview",        "DoctorName, Specialization, City, Rating, IsApproved", "Модерирано от общността"],
            ["ChildFriendlyPlace",  "Name, PlaceType, City, AgeFromMonths, AgeToMonths, IsApproved", "Общностна директория"],
            ["SavedSearch",         "UserId, Name, Filters (категория, вид, размер, цена…)", "Задейства известие при нови артикули"],
            ["AuditLog",            "UserId, Action, EntityType, EntityId, Changes, CreatedAt", "История на административни действия"],
        ],
        col_widths=[1.7, 3.0, 1.9]
    )

    heading2(doc, "5.4  Структура на фронтенда")
    add_table(doc,
        headers=["Директория", "Съдържание"],
        rows=[
            ["src/api/",        "Axios клиентски модули по домейн (auth, items, payments, shipping, messages…)"],
            ["src/store/",      "Zustand хранилища: authStore, itemStore, chatStore, notificationStore…"],
            ["src/contexts/",   "React контексти: SignalR връзка, auth доставчик"],
            ["src/pages/",      "Един файл на route: HomePage, BrowsePage, ItemDetailPage, DashboardPage, ChatPage, AdminPage…"],
            ["src/components/", "Многократно използваеми UI компоненти: Navbar, ItemCard, PaymentModal, ShipmentTracker, ChatWindow…"],
            ["src/layouts/",    "MainLayout (навигация + долен колонтитул), AuthLayout (центриран), AdminLayout (странично меню)"],
            ["src/locales/",    "Файлове с преводи en.json и bg.json"],
            ["src/hooks/",      "Персонализирани hooks: useAuth, useSignalR, useInfiniteScroll…"],
        ],
        col_widths=[1.8, 4.8]
    )


def section_api_reference(doc):
    heading1(doc, "6. Справочник за API")

    body(doc,
        "Всички endpoints са версионирани под /api/v1/. Защитените endpoints изискват "
        "валиден JWT Bearer токен. Admin endpoints изискват роля Admin."
    )

    heading2(doc, "6.1  Автентикация  (/api/v1/auth)")
    add_table(doc,
        headers=["Метод", "Път", "Достъп", "Описание"],
        rows=[
            ["POST", "/register",        "Публичен",     "Регистрация с имейл, парола и Turnstile токен"],
            ["POST", "/login",           "Публичен",     "Вход с имейл/парола; връща access + refresh токен"],
            ["POST", "/refresh",         "Публичен",     "Обменя refresh токен за нов access токен"],
            ["POST", "/google-login",    "Публичен",     "Вход / регистрация с Google OAuth"],
            ["GET",  "/me",              "Оторизиран",   "Взема профила на текущия потребител"],
            ["POST", "/revoke",          "Оторизиран",   "Анулира текущия refresh токен (изход)"],
            ["POST", "/change-password", "Оторизиран",   "Смяна на парола (изисква текущата парола)"],
            ["POST", "/forgot-password", "Публичен",     "Изпраща имейл за възстановяване на парола"],
            ["POST", "/reset-password",  "Публичен",     "Завършва нулирането с токен + нова парола"],
        ],
        col_widths=[0.8, 2.0, 1.2, 2.6]
    )

    heading2(doc, "6.2  Обяви  (/api/v1/items)")
    add_table(doc,
        headers=["Метод", "Път", "Достъп", "Описание"],
        rows=[
            ["GET",   "/",           "Публичен",      "Разглеждане на обяви (страниране, филтри, сортиране). Кеширано 30с."],
            ["POST",  "/",           "Оторизиран",    "Създаване на нова обява"],
            ["GET",   "/{id}",       "Публичен",      "Детайли за обява; увеличава брояча на прегледи"],
            ["PATCH", "/{id}",       "Собственик",    "Актуализиране на полета на обявата"],
            ["DELETE","/{id}",       "Собственик",    "Изтриване на обявата"],
            ["POST",  "/{id}/like",  "Оторизиран",    "Превключване харесване/отхаресване на обява"],
            ["POST",  "/{id}/bump",  "Собственик",    "Освежаване на обявата (нулиране на createdAt)"],
            ["POST",  "/ai-suggest", "Оторизиран",    "Изпращане на снимка → AI предложения за обява"],
        ],
        col_widths=[0.8, 1.7, 1.2, 3.0]
    )

    heading2(doc, "6.3  Плащания  (/api/v1/payments)")
    add_table(doc,
        headers=["Метод", "Път", "Достъп", "Описание"],
        rows=[
            ["POST", "/checkout/{itemId}",   "Оторизиран", "Създаване на Stripe Checkout сесия за картово плащане"],
            ["POST", "/on-spot/{itemId}",    "Оторизиран", "Записване на плащане на място (в брой)"],
            ["POST", "/webhook",             "Публичен*",  "Stripe webhook обработчик (проверка на Stripe-Signature)"],
            ["GET",  "/",                    "Оторизиран", "Списък с плащанията на текущия потребител"],
            ["POST", "/wallet/topup",        "Оторизиран", "Създаване на PaymentIntent за зареждане на портфейл"],
        ],
        col_widths=[0.8, 2.2, 1.2, 2.5]
    )

    heading2(doc, "6.4  Доставка  (/api/v1/shipping)")
    add_table(doc,
        headers=["Метод", "Път", "Достъп", "Описание"],
        rows=[
            ["POST", "/calculate",             "Оторизиран", "Изчисляване на цена за избран куриер и вид доставка"],
            ["POST", "/create",                "Оторизиран", "Създаване на товарителница с избрания куриер"],
            ["GET",  "/offices/{provider}",    "Оторизиран", "Списък с офиси/автомати за даден куриер"],
            ["GET",  "/{shipmentId}",          "Оторизиран", "Детайли за пратка и текущ статус"],
            ["GET",  "/label/{shipmentId}",    "Оторизиран", "Изтегляне на PDF етикет за доставка"],
            ["POST", "/{shipmentId}/track",    "Оторизиран", "Синхронизиране на статуса от куриерското API"],
            ["POST", "/{shipmentId}/cancel",   "Оторизиран", "Анулиране на активна пратка"],
        ],
        col_widths=[0.8, 2.2, 1.2, 2.5]
    )

    heading2(doc, "6.5  Допълнителни endpoints")
    add_table(doc,
        headers=["Контролер", "Базов път", "Основни endpoints"],
        rows=[
            ["Съобщения",           "/api/v1/messages",               "GET разговори, GET /userId (история), POST /mark-read"],
            ["Е-Фактури",           "/api/v1/ebills",                 "GET списък, POST /{id}/resend (с ограничаване)"],
            ["Пакети",              "/api/v1/bundles",                "CRUD + заявки за покупка"],
            ["Потребители",         "/api/v1/users",                  "GET профил, PATCH себе си, GET /{id}/items, GET /{id}/ratings"],
            ["Категории",           "/api/v1/categories",             "GET списък (кеширан 1ч)"],
            ["Оферти",              "/api/v1/offers",                 "POST създаване, PATCH контра/приемане/отхвърляне"],
            ["Заявки за покупка",   "/api/v1/purchase-requests",      "POST създаване, POST приемане/отхвърляне"],
            ["Следване",            "/api/v1/follows",                "POST следване, DELETE прекратяване, GET последователи"],
            ["Оценки",              "/api/v1/ratings",                "POST оценяване на потребител след покупка"],
            ["Рецензии за лекари",  "/api/v1/doctor-reviews",         "POST изпращане, GET одобрен списък (admin: одобряване/отхвърляне)"],
            ["Детски места",        "/api/v1/child-friendly-places",  "POST изпращане, GET филтриран списък (admin: одобряване/отхвърляне)"],
            ["Запазени търсения",   "/api/v1/saved-searches",         "POST създаване, GET списък, DELETE премахване"],
            ["Обратна връзка",      "/api/v1/feedback",               "POST изпращане, admin: GET списък"],
            ["Снимки",              "/api/v1/photos",                 "POST качване (R2), DELETE премахване"],
            ["Admin",               "/api/v1/admin",                  "Статистики на таблото, блокиране на потребители, модерация, одитен журнал"],
            ["Асистент",            "/api/v1/assistant",              "POST чат съобщение, GET история на разговора"],
        ],
        col_widths=[1.9, 2.2, 2.5]
    )


def section_automation(doc):
    heading1(doc, "7. Автоматизация и работни процеси (n8n)")

    body(doc,
        "MamVibe се доставя с 16 предварително изградени n8n workflow файла в директорията "
        "n8n-workflows/. Те се импортират в самостоятелно хостван n8n инстанс и се задействат "
        "чрез HTTP webhooks, изпращани от фоновата услуга N8nWebhookService на платформата."
    )

    add_table(doc,
        headers=["Събитие", "Задействащо условие", "Автоматизирано действие"],
        rows=[
            ["payment.completed",        "Stripe checkout потвърден",           "Имейл за потвърждение на покупка до купувача и продавача"],
            ["shipment.created",         "Нова товарителница създадена",         "Имейл до купувача с номер за проследяване и данни за куриера"],
            ["shipment.delivered",       "Статус на куриера → Доставена",       "Имейл и до двете страни; напомняне до купувача за оценка"],
            ["shipment.stuck",           "Ежедневно: в транзит 7+ дни",         "Предупреждение до администратора; известие до купувача с линк за проследяване"],
            ["user.registered",          "Нов акаунт създаден",                 "Приветствен имейл с ръководство за платформата"],
            ["user.blocked",             "Администратор блокира акаунт",        "Имейл за уведомление до засегнатия потребител"],
            ["chat.message_offline",     "Съобщение до офлайн потребител",      "Имейл с откъс от съобщението и линк за отговор"],
            ["stale_items",             "Ежедневно: публикувани 30+ дни",       "Напомняне до продавача да прегледа, намали или премахне"],
            ["daily_summary",            "Всеки ден в 08:00 UTC",              "Административен отчет: нови обяви, плащания, пратки"],
            ["feedback_prompt",          "Ежедневно: доставени 2+ дни, без оценка", "Напомняне до купувача за оценка на продавача"],
            ["weekly_seller_report",     "Понеделник в 09:00 ч.",              "Имейл до продавача: обобщение на продажби, приход, средна оценка"],
        ],
        col_widths=[2.0, 2.0, 2.6]
    )


def section_security(doc):
    heading1(doc, "8. Сигурност и съответствие")

    heading2(doc, "8.1  Мерки за сигурност")
    add_table(doc,
        headers=["Мярка", "Имплементация"],
        rows=[
            ["Сигурност на транспорта",    "HTTPS задължително; HSTS заглавие (max-age=31536000)"],
            ["API автентикация",           "JWT Bearer токени; кратък живот на access токена; HttpOnly refresh бисквитки"],
            ["OAuth",                      "Google OAuth 2.0 чрез ASP.NET Identity external login"],
            ["CAPTCHA",                    "Cloudflare Turnstile при регистрация, вход и нулиране на парола"],
            ["Политика за пароли",         "Мин. 8 знака, главна буква, малка буква, цифра, специален знак (ASP.NET Identity)"],
            ["Блокиране при грешни опити", "5 неуспешни влизания → 5-минутно блокиране"],
            ["Ограничаване на заявките",   "Глобално 200 заявки/мин; автентикация 30/мин; качване 20/мин; е-фактура 3/мин"],
            ["Валидация на входа",         "FluentValidation за всички DTO заявки; без директно свързване на entities"],
            ["CORS",                       "Явен списък с доверени произходи; без wildcard"],
            ["HTTP заглавки за сигурност", "CSP, X-Frame-Options (DENY), X-Content-Type-Options, Referrer-Policy, Permissions-Policy, CORP, COOP"],
            ["Блокирани потребители",      "Middleware отхвърля всички заявки от блокирани акаунти при всяко извикване"],
            ["Защита на метриките",        "/metrics endpoint връща 404 за не-вътрешни IP адреси"],
            ["Stripe webhook валидация",   "Заглавието Stripe-Signature се проверява преди всяка обработка на плащане"],
            ["Data Protection ключове",    "ASP.NET Data Protection ключовете са персистирани в Docker volume"],
            ["Чувствителни данни в логове","Само entity ID-та и статус кодове; без лични данни или финансова информация"],
            ["GDPR",                       "Endpoint за експорт на данни — потребителите могат да изтеглят всичките си лични данни"],
            ["Swagger",                    "API документацията се обслужва само в Development среда"],
        ],
        col_widths=[2.2, 4.4]
    )

    heading2(doc, "8.2  Сигурност на зависимостите")
    body(doc,
        "Всички NuGet пакети се проверяват редовно с командата dotnet list package --vulnerable. "
        "В дървото на зависимостите няма известни уязвими пакети. Frontend npm пакетите "
        "се проверяват чрез npm audit."
    )


def section_localisation(doc):
    heading1(doc, "9. Локализация и достъпност")

    heading2(doc, "9.1  Езикова поддръжка")
    body(doc,
        "MamVibe поддържа английски и български на целия фронтенд. Определянето на езика "
        "следва следния приоритет: предпочитание, съхранено в акаунта → localStorage → "
        "езикът на браузъра → резервен вариант — английски."
    )
    add_table(doc,
        headers=["Локал", "Покритие"],
        rows=[
            ["Английски (en)", "Пълен превод на UI — всички етикети, съобщения и низове за грешки"],
            ["Български (bg)", "Пълен превод на UI — всички етикети, съобщения и низове за грешки"],
        ],
        col_widths=[1.5, 5.1]
    )

    heading2(doc, "9.2  Фискална и правна локализация")
    for f in [
        "Прилагана ставка на ДДС от 20% (стандартна за България) към всички облагаеми транзакции.",
        "Интеграция с TakeANap за фискален бон за НАП.",
        "Номерацията на е-фактурите следва местните конвенции.",
        "Форматирането на дати и числа се адаптира спрямо избрания locale чрез date-fns.",
        "Страници с Политика за поверителност, Общи условия и Политика за бисквитки на двата езика.",
    ]:
        bullet(doc, f)


def section_deployment(doc):
    heading1(doc, "10. Разгръщане и операции")

    heading2(doc, "10.1  Docker Compose стек")
    add_table(doc,
        headers=["Услуга", "Образ / Технология", "Ресурси", "Роля"],
        rows=[
            ["postgres",    "PostgreSQL 18",         "512 MB RAM, 1 CPU", "Основна база данни с healthcheck"],
            ["pgbouncer",   "PgBouncer",             "Лек",               "Connection pooler (transaction mode, 500 клиента)"],
            ["redis",       "Redis 7",               "256 MB LRU + AOF",  "Кеш + SignalR backplane"],
            ["api",         ".NET 8 / ASP.NET Core", "1 GB RAM, 2 CPU",   "Backend REST API + SignalR хъб"],
            ["frontend",    "Nginx + React build",   "Лек",               "Статични файлове + обратен прокси"],
            ["prometheus",  "Prometheus",            "256 MB",            "Събирач на метрики"],
            ["loki",        "Grafana Loki",          "256 MB",            "Агрегиране на логове (31-дневно задържане)"],
            ["grafana",     "Grafana",               "256 MB",            "Табла за мониторинг и предупреждения"],
        ],
        col_widths=[1.3, 1.9, 1.5, 2.0]
    )

    heading2(doc, "10.2  Команди за бърз старт")
    heading3(doc, "Пълен стек (Docker)")
    for cmd in [
        "docker compose up --build    # Изграждане и стартиране на всички услуги",
        "docker compose down          # Спиране на всички услуги",
    ]:
        bullet(doc, cmd)

    heading3(doc, "Режим за разработка (локални услуги)")
    for cmd in [
        "docker compose up postgres -d                     # Само базата данни",
        "dotnet run  (от backend/src/MomVibe.WebApi/)      # API на localhost:5038",
        "npm run dev (от frontend/)                        # UI на localhost:5173",
    ]:
        bullet(doc, cmd)

    heading2(doc, "10.3  Променливи на средата")
    body(doc,
        "Всички тайни се подават чрез файл .env в корена (вижте .env.example за пълния шаблон). "
        "Doppler интеграцията е налична за управление на тайни в продукция."
    )
    add_table(doc,
        headers=["Група", "Ключови променливи"],
        rows=[
            ["PostgreSQL",        "POSTGRES_DB, POSTGRES_USER, POSTGRES_PASSWORD, POSTGRES_PORT"],
            ["JWT",               "JWT_SECRET (минимум 32 символа)"],
            ["OAuth",             "GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET, FRONTEND_URL"],
            ["Stripe",            "STRIPE_SECRET_KEY, STRIPE_WEBHOOK_SECRET, STRIPE_WALLET_WEBHOOK_SECRET"],
            ["Куриери",           "ECONT_USERNAME/PASSWORD, SPEEDY_USERNAME/PASSWORD, BOXNOW_API_KEY, PIGEONEXPRESS_API_KEY"],
            ["Фискален бон",      "TAKEANAP_API_KEY, TAKEANAP_API_SECRET, TAKEANAP_SHOP_ID"],
            ["Имейл / SMTP",      "SMTP_HOST, SMTP_PORT, SMTP_USERNAME, SMTP_PASSWORD, SMTP_FROM_EMAIL"],
            ["AI",                "ANTHROPIC_API_KEY, GROQ_API_KEY, AI_CHAT_PROVIDER (anthropic|groq)"],
            ["Cloudflare",        "TURNSTILE_SITE_KEY, TURNSTILE_SECRET_KEY, R2 credentials"],
            ["n8n",               "N8N_BASE_URL, N8N_ENABLED, N8N_WEBHOOK_SECRET"],
            ["Наблюдаемост",      "OpenTelemetry__Otlp__Endpoint"],
        ],
        col_widths=[1.7, 4.9]
    )

    heading2(doc, "10.4  Стратегия за клоновете (GitFlow)")
    add_table(doc,
        headers=["Клон", "Предназначение"],
        rows=[
            ["main",          "Винаги готов за разгръщане. Само тестван, продукционно готов код."],
            ["develop",       "Интеграционен клон. Feature клоновете се сливат тук първо."],
            ["feature/*",     "Клон от develop; един клон на функционалност или поправка на грешка."],
            ["release/x.y.z", "Стабилизация преди продукция. Слива се в main и develop."],
            ["hotfix/*",      "Спешна поправка в продукция. Клон от main; слива се в main и develop."],
        ],
        col_widths=[1.6, 5.0]
    )


def section_team_guide(doc):
    heading1(doc, "11. Ръководство за въвеждане на разработчици")

    heading2(doc, "11.1  Структура на хранилището")
    add_table(doc,
        headers=["Път", "Съдържание"],
        rows=[
            ["backend/src/MomVibe.Domain/",         "Entities, enums, домейн константи"],
            ["backend/src/MomVibe.Application/",    "DTO, интерфейси, валидатори, AutoMapper профили"],
            ["backend/src/MomVibe.Infrastructure/", "EF Core, миграции, клиенти на външни услуги"],
            ["backend/src/MomVibe.WebApi/",         "Контролери, SignalR хъбове, middleware, StartUp.cs"],
            ["backend/tests/MomVibe.UnitTests/",    "xUnit unit тестове (с mock зависимости)"],
            ["backend/tests/MomVibe.IntegrationTests/", "xUnit интеграционни тестове (реална DB)"],
            ["frontend/src/",                       "React 19 изходен код"],
            ["n8n-workflows/",                      "16 предварително изградени n8n workflow JSON файла"],
            ["docs/",                               "ADR-и, скрийншоти, тази документация"],
        ],
        col_widths=[3.0, 3.6]
    )

    heading2(doc, "11.2  Чести команди")
    heading3(doc, "Бекенд")
    for cmd in [
        "dotnet run                                              # Стартиране на API (от MomVibe.WebApi/)",
        "dotnet test backend/tests/MomVibe.UnitTests            # Изпълнение на unit тестове",
        "dotnet test backend/tests/MomVibe.IntegrationTests     # Изпълнение на интеграционни тестове",
        "dotnet ef migrations add <Ime> --project ../MomVibe.Infrastructure --startup-project .",
        "dotnet ef database update --project ../MomVibe.Infrastructure --startup-project .",
    ]:
        bullet(doc, cmd)

    heading3(doc, "Фронтенд")
    for cmd in [
        "npm run dev      # Стартиране на Vite dev сървър на localhost:5173",
        "npm run build    # TypeScript проверка + продукционен bundle",
        "npm run lint     # ESLint",
        "npm run preview  # Локален преглед на продукционния build",
    ]:
        bullet(doc, cmd)

    heading2(doc, "11.3  Бележка за именуването")
    body(doc,
        "Хранилището и Docker Compose файлът използват името MamVibe. "
        "Въпреки това, всички .NET имена на проекти, namespaces и C# изходен код използват MomVibe "
        "(напр. MomVibe.Domain, MomVibe.Application). Новият код трябва да следва "
        "съществуващата конвенция: namespace MomVibe.*"
    )

    heading2(doc, "11.4  Чеклист при добавяне на нова функционалност")
    for step in [
        "Създайте или актуализирайте домейн entities в MomVibe.Domain/.",
        "Добавете или актуализирайте интерфейса в MomVibe.Application/Interfaces/.",
        "Добавете DTO заявката/отговора и FluentValidation валидатора в MomVibe.Application/.",
        "Имплементирайте услугата и всички клиенти на външни API в MomVibe.Infrastructure/.",
        "Регистрирайте услугите в съответния метод AddInfrastructureServices().",
        "Добавете контролер endpoint в MomVibe.WebApi/Controllers/.",
        "Напишете unit тестове в MomVibe.UnitTests/ и интеграционни в MomVibe.IntegrationTests/.",
        "Добавете или актуализирайте frontend API клиент в frontend/src/api/.",
        "Актуализирайте Zustand store, ако е нужно глобално състояние.",
        "Добавете страницата/компонента в frontend/src/pages/ или frontend/src/components/.",
        "Добавете ключовете за превод в двата файла: en.json и bg.json.",
        "Добавете нова EF Core миграция, ако схемата е променена.",
    ]:
        bullet(doc, step)


def section_roadmap(doc):
    heading1(doc, "12. Текущо състояние и пътна карта")

    heading2(doc, "12.1  Завършено (готово за продукция)")
    for item in [
        "Пълен пазар: създаване, разглеждане, покупка и продажба на обяви.",
        "Видове обяви за дарение и продажба с AI фото асистент.",
        "AI модерация на съдържание с оценки на увереност.",
        "Пълен стек за плащания: Stripe, Портфейл, На място, Наложен платеж.",
        "Пълна интеграция с доставчици: Еконт, Speedy, Box Now, Pigeon Express.",
        "SignalR чат в реално време с индикатори за писане и потвърждения за прочитане.",
        "AI чат бот (Anthropic Claude / Groq).",
        "Е-фактури и TakeANap фискален бон.",
        "Административен панел: потребители, обяви, пратки, одитен журнал.",
        "Рецензии за лекари и детски места — общностни директории.",
        "Система за следване/последователи и оценки след покупка.",
        "Запазени търсения с имейл известия.",
        "n8n автоматизация: 16 предварително изградени workflow-и.",
        "Двуезичен UI (английски и български).",
        "Пълен стек за наблюдаемост (Prometheus, Loki, Grafana).",
        "Cloudflare Turnstile защита от ботове.",
        "Експорт на данни по GDPR.",
        "Разгръщане в Docker Compose за продукция.",
    ]:
        bullet(doc, item)

    heading2(doc, "12.2  Потенциални следващи стъпки")
    for item in [
        "Нативно мобилно приложение (React Native / Expo) с пълна поддръжка на push известия.",
        "Разширени интеграции с куриери за допълнителни европейски пазари.",
        "Абонаментни нива за продавачи (featured обяви, приоритетна поддръжка).",
        "Разширено търсене: разглеждане на обяви на карта, филтриране по радиус.",
        "Геймификация: постижения за продавачи, значки, програма за лоялност.",
        "Разширени AI функции: автоматично сравняване на цени, оценки на качеството на обявите.",
        "Локализация за нови страни (гръцки, румънски) за разширяване на Балканите.",
    ]:
        bullet(doc, item)


def section_glossary(doc):
    heading1(doc, "13. Речник")

    add_table(doc,
        headers=["Термин", "Определение"],
        rows=[
            ["MamVibe",         "Брандовото и хранилищното наименование на платформата."],
            ["MomVibe",         "Namespace, използван в целия .NET код (MomVibe.Domain и др.)"],
            ["Item (Артикул)",  "Единичен продукт, публикуван за дарение или продажба."],
            ["Bundle (Пакет)",  "Група от 2-10 артикула, продавани заедно на намалена обща цена."],
            ["Listing Type",    "Дарение (безплатно) или Продажба (с цена)."],
            ["Purchase Request","Официална заявка на купувача за придобиване на артикул; продавачът трябва да я приеме."],
            ["Offer (Оферта)",  "Ценово предложение под обявената цена; продавачът може да направи контраоферта."],
            ["Wallet (Портфейл)","Вътрешен ескроу акаунт, задържащ средствата на купувача до потвърждаване на доставката."],
            ["E-Bill (Е-Фактура)","Електронна разписка за плащане, генерирана автоматично за всяка завършена покупка."],
            ["Waybill (Товарит.)","Куриерски транспортен документ; уникална референция за проследяване на пратка."],
            ["TakeANap",        "Български доставчик на фискален бон, интегриран за съответствие с ДДС."],
            ["n8n",             "Open-source инструмент за автоматизация на работни процеси за имейли и известия."],
            ["Turnstile",       "Невидимият CAPTCHA услуга на Cloudflare при автентикация."],
            ["PgBouncer",       "PostgreSQL connection pooler, работещ в transaction mode."],
            ["SignalR",         "Библиотеката на Microsoft за WebSocket, използвана за чат и известия в реално време."],
            ["R2",              "S3-съвместимото object storage на Cloudflare за качване на снимки."],
            ["Groq",            "AI доставчик с ниска латентност, използван като алтернатива на Anthropic."],
            ["GitFlow",         "Стратегия за клонове: main, develop, feature/*, release/*, hotfix/*."],
            ["Clean Architecture","Слоеста структура на кода: Domain → Application → Infrastructure → WebApi."],
            ["COD",             "Наложен платеж — куриерът събира парите на вратата на получателя."],
        ],
        col_widths=[1.8, 4.8]
    )


# ──────────────────────────────────────────────────────────
# Основен builder
# ──────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    for section in doc.sections:
        section.page_width    = Inches(8.5)
        section.page_height   = Inches(11)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = WARM_GRAY

    # Корица
    build_cover(doc)

    # Съдържание
    heading1(doc, "Съдържание")
    toc_items = [
        ("1.",  "Резюме"),
        ("2.",  "Преглед на продукта"),
        ("3.",  "Функционалности"),
        ("4.",  "Технологичен стек"),
        ("5.",  "Архитектура"),
        ("6.",  "Справочник за API"),
        ("7.",  "Автоматизация и работни процеси (n8n)"),
        ("8.",  "Сигурност и съответствие"),
        ("9.",  "Локализация и достъпност"),
        ("10.", "Разгръщане и операции"),
        ("11.", "Ръководство за въвеждане на разработчици"),
        ("12.", "Текущо състояние и пътна карта"),
        ("13.", "Речник"),
    ]
    for num, title in toc_items:
        p    = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        rnum = p.add_run(f"{num}  ")
        rnum.bold = True
        rnum.font.size = Pt(10.5)
        rnum.font.color.rgb = PURPLE
        rtitle = p.add_run(title)
        rtitle.font.size = Pt(10.5)
        rtitle.font.color.rgb = WARM_GRAY

    page_break(doc)

    section_executive_summary(doc)
    page_break(doc)

    section_product_overview(doc)
    page_break(doc)

    section_features(doc)
    page_break(doc)

    section_tech_stack(doc)
    page_break(doc)

    section_architecture(doc)
    page_break(doc)

    section_api_reference(doc)
    page_break(doc)

    section_automation(doc)
    section_security(doc)
    page_break(doc)

    section_localisation(doc)
    section_deployment(doc)
    page_break(doc)

    section_team_guide(doc)
    page_break(doc)

    section_roadmap(doc)
    section_glossary(doc)

    # Долен колонтитул
    page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 60)
    run.font.color.rgb = PURPLE_LIGHT
    run.font.size = Pt(10)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(
        f"MamVibe Преглед на платформата  ·  Поверително  ·  {datetime.date.today().year}"
    )
    r2.font.size = Pt(9)
    r2.font.color.rgb = WARM_GRAY
    r2.italic = True

    return doc


if __name__ == "__main__":
    output_path = r"C:\WORK_PLACE\MamVibe\docs\word_document\MamVibe_Преглед_на_платформата_BG.docx"
    doc = build_document()
    doc.save(output_path)
    print("Document saved: " + output_path)
